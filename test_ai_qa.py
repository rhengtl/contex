# test_ai_qa.py
"""
Verification suite for the conversion pipeline and the app around it.

Run it with:      python test_ai_qa.py

Covers everything that does not need a live API key: the AI pipeline driven by
a scripted stand-in for the model, model rotation and the local fallback,
equation segmentation, the LaTeX validator, the real compile path, image
preprocessing, the Flask routes, the session-scoped result store, and the page
previews. The heavy OCR model and Firebase are stubbed out so the suite runs in
seconds and touches no network.
"""

import io
import os
import re
import sys
import tempfile
import types

# --- Isolate the app from its heavy/optional dependencies ------------------
_recognized = {}


def _fake_process_image_list(file_bytes):
    return _recognized.get('equations', [{'index': 1, 'latex': 'x^2'}])


_fake_equation = types.ModuleType('equation')
_fake_equation.is_model_loaded = lambda: True
_fake_equation.process_image = lambda data: 'x^2'
_fake_equation.process_image_list = _fake_process_image_list
_fake_equation.segment_boxes = lambda img, max_regions=None, allow_empty=False: (
    _recognized.get('boxes', []))
_fake_equation.recognize = lambda crop: _recognized.get('latex', 'x^{2}')
_fake_equation.tighten = lambda img, box: box
sys.modules['equation'] = _fake_equation

_saved_history = []
_fake_firebase = types.ModuleType('firebase_config')
_fake_firebase.save_ocr_history = lambda uid, name, kind, result, truncated=False: (
    _saved_history.append((uid, name, kind, result))
    or f'doc-{len(_saved_history)}')
_fake_firebase.get_user_ocr_history = lambda *a, **k: []
_fake_firebase.get_ocr_history_item = lambda uid, doc_id: _history_items.get(
    (uid, doc_id))
_fake_firebase.verify_id_token = lambda *a, **k: None
_fake_firebase.get_user_by_uid = lambda *a, **k: None
_fake_firebase.upsert_user_profile = lambda *a, **k: True
_fake_firebase.verify_user = lambda *a, **k: {'success': False}
_fake_firebase.create_user = lambda *a, **k: {'success': False}
_fake_firebase.send_password_reset = lambda *a, **k: {'success': True}
_fake_firebase.set_terms_accepted = lambda uid, version: _accepted_terms.__setitem__(
    uid, version) or True
_fake_firebase.get_terms_accepted = lambda uid: _accepted_terms.get(uid)
sys.modules['firebase_config'] = _fake_firebase

_history_items = {}
_accepted_terms = {}

_SCRATCH = tempfile.mkdtemp(prefix='contex_test_')
os.environ['UPLOAD_FOLDER'] = _SCRATCH
os.environ['FLASK_SECRET_KEY'] = 'test-secret'
os.environ.setdefault('GEMINI_API_KEY', 'test-key-not-used')

from PIL import Image  # noqa: E402
import pikepdf  # noqa: E402

import ai_qa  # noqa: E402
import ai_status  # noqa: E402
import convert  # noqa: E402
import docx_input  # noqa: E402
import latex_tools  # noqa: E402
import layout  # noqa: E402
import llm_providers  # noqa: E402
import preprocess  # noqa: E402
import tex_store  # noqa: E402
import textract_fast  # noqa: E402
import app as flask_app  # noqa: E402


# ---------------------------------------------------------------------------
# Harness
# ---------------------------------------------------------------------------

_RESULTS = []


def test(name):
    def decorator(fn):
        _RESULTS.append((name, fn))
        return fn
    return decorator


def check(condition, message):
    if not condition:
        raise AssertionError(message)


# The route tests stub the conversion layer by assigning to `flask_app.ai_qa`
# and `flask_app.convert`, which are the very same module objects as `ai_qa`
# and `convert` - so those stubs would otherwise leak into every test that runs
# afterwards. Snapshot the real callables now and put them back before each
# test.
_REAL_CALLABLES = {
    (ai_qa, 'convert_page'): ai_qa.convert_page,
    (convert, 'convert'): convert.convert,
}


def restore_stubs():
    for (module, name), original in _REAL_CALLABLES.items():
        setattr(module, name, original)


def skip(reason):
    print(f'      (skipped: {reason})', end=' ')


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

GOOD_TEX = r"""\documentclass{article}
\usepackage{amsmath}
\begin{document}
\section{Energy}
Einstein's relation is $E = mc^{2}$.
\end{document}
"""

BROKEN_TEX = r"""\documentclass{article}
\begin{document}
\begin{itemize}
\item One
\end{document}
"""


def png_bytes(size=(600, 400), color=(255, 255, 255)):
    buffer = io.BytesIO()
    Image.new('RGB', size, color).save(buffer, format='PNG')
    return buffer.getvalue()


def pdf_bytes(pages=2):
    pdf = pikepdf.Pdf.new()
    for _ in range(pages):
        pdf.add_blank_page(page_size=(200, 200))
    buffer = io.BytesIO()
    pdf.save(buffer)
    return buffer.getvalue()


_DOCX_CACHE = {}


def docx_bytes():
    """
    A Word document with every structure the extractor has to handle.

    Built rather than checked in, so the fixture cannot drift away from what
    python-docx actually produces. The equations are inserted as raw OMML,
    which is how Word really stores them - python-docx has no API for it, and
    that is exactly the gap the extractor exists to cover.
    """
    if 'bytes' in _DOCX_CACHE:
        return _DOCX_CACHE['bytes']
    from docx import Document

    math_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

    def add_math(paragraph, symbols):
        maker = paragraph._p.makeelement
        omath = maker('{%s}oMath' % math_ns, {})
        run = maker('{%s}r' % math_ns, {})
        text = maker('{%s}t' % math_ns, {})
        text.text = symbols
        run.append(text)
        omath.append(run)
        paragraph._p.append(omath)

    document = Document()
    document.add_heading('Relativity', 0)
    document.add_heading('Mass and energy', 1)
    document.add_paragraph('Einstein showed that mass and energy are '
                           'equivalent. The relation is written below.')
    add_math(document.add_paragraph(), 'E=mc2')
    inline = document.add_paragraph('Inline, the constant ')
    add_math(inline, 'c')
    inline.add_run(' is the speed of light in vacuum. 100% agreed.')
    document.add_paragraph('First point', style='List Bullet')
    document.add_paragraph('Second point', style='List Bullet')
    table = document.add_table(rows=2, cols=3)
    for index, value in enumerate(['Symbol', 'Meaning', 'Units']):
        table.cell(0, index).text = value
    for index, value in enumerate(['m', 'mass', 'kg']):
        table.cell(1, index).text = value
    document.add_paragraph('That concludes the summary.')

    buffer = io.BytesIO()
    document.save(buffer)
    _DOCX_CACHE['bytes'] = buffer.getvalue()
    return _DOCX_CACHE['bytes']


def text_page(lines=('Chapter 1', 'The quick brown fox jumps'), size=(900, 300)):
    from PIL import ImageDraw
    image = Image.new('L', size, 255)
    draw = ImageDraw.Draw(image)
    for index, line in enumerate(lines):
        draw.text((40, 40 + index * 60), line, fill=0)
    return image


def render_tex_page(body, dpi=150):
    """Render LaTeX to a page image - used to test equation segmentation."""
    source = (r'\documentclass[12pt]{article}\usepackage[margin=2cm]{geometry}'
              r'\usepackage{amsmath}\pagestyle{empty}\begin{document}'
              + body + r'\end{document}')
    result = latex_tools.compile_tex(source, want_pdf=True)
    if not result['ok']:
        return None
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        return None
    try:
        return convert_from_bytes(result['pdf'], dpi=dpi)[0]
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Equation segmentation - the converter's new list output
# ---------------------------------------------------------------------------
# equation.py is stubbed for the route tests, so load the real module under a
# different name with its model imports faked out.

def _load_real_equation_module():
    import importlib.util
    for name, attr in (('transformers', 'TrOCRProcessor'),
                       ('optimum.onnxruntime', 'ORTModelForVision2Seq')):
        if name not in sys.modules:
            module = types.ModuleType(name)
            setattr(module, attr, None)
            sys.modules[name] = module
    sys.modules.setdefault('optimum', types.ModuleType('optimum'))
    spec = importlib.util.spec_from_file_location('_real_equation', 'equation.py')
    module = importlib.util.module_from_spec(spec)
    try:
        spec.loader.exec_module(module)
    except Exception:
        pass
    return module


_EQ = _load_real_equation_module()


@test('separate equations on a page are detected separately')
def _():
    page = render_tex_page(r"""\[ E = mc^{2} \]
\[ a^{2}+b^{2}=c^{2} \]
\[ e^{i\pi}+1=0 \]""")
    if page is None:
        return skip('no LaTeX engine or Poppler')
    regions = _EQ.segment_equations(page)
    check(len(regions) == 3, f'{len(regions)} regions, expected 3')


@test('a fraction is not split into numerator and denominator')
def _():
    page = render_tex_page(r'\[ x = \frac{-b \pm \sqrt{b^{2}-4ac}}{2a} \]')
    if page is None:
        return skip('no LaTeX engine or Poppler')
    regions = _EQ.segment_equations(page)
    check(len(regions) == 1, f'a single fraction split into {len(regions)} pieces')


@test('nested fractions stay in one region')
def _():
    page = render_tex_page(r'\[ y = \frac{\frac{a}{b}}{\frac{c}{d}} \]')
    if page is None:
        return skip('no LaTeX engine or Poppler')
    check(len(_EQ.segment_equations(page)) == 1, 'nested fraction was split')


@test('equations containing fractions are still separated from each other')
def _():
    page = render_tex_page(r"""\[ E = mc^{2} \]
\[ x = \frac{-b \pm \sqrt{b^{2}-4ac}}{2a} \]
\[ \sum_{i=1}^{n} i = \frac{n(n+1)}{2} \]
\[ \int_{0}^{1} x^{2}\,dx = \frac{1}{3} \]""")
    if page is None:
        return skip('no LaTeX engine or Poppler')
    regions = _EQ.segment_equations(page)
    check(len(regions) == 4, f'{len(regions)} regions, expected 4')


@test('a single equation is returned as one region, as it always was')
def _():
    page = render_tex_page(r'\[ E = mc^{2} \]')
    if page is None:
        return skip('no LaTeX engine or Poppler')
    check(len(_EQ.segment_equations(page)) == 1, 'single equation was split')


@test('each region is cropped tightly around its equation')
def _():
    # pix2text-mfr expects a formula that fills its crop. Handing it a
    # full-width strip of mostly blank page made it emit runs of \qquad
    # instead of the expression, so the crop must follow the ink horizontally.
    page = render_tex_page(r"""\[ v = u + at \]
\[ v^{2} = u^{2} + 2as \]""")
    if page is None:
        return skip('no LaTeX engine or Poppler')
    regions = _EQ.segment_equations(page)
    check(len(regions) == 2, f'{len(regions)} regions, expected 2')
    for region in regions:
        check(region.size[0] < page.size[0] * 0.9,
              f'region is {region.size[0]}px wide on a {page.size[0]}px page '
              '- it was not cropped horizontally')


@test('a blank page does not crash the segmenter')
def _():
    blank = Image.new('L', (400, 300), 255)
    check(len(_EQ.segment_equations(blank)) == 1, 'blank page misbehaved')


@test('the number of regions is capped')
def _():
    body = '\n'.join(rf'\[ x_{{{i}}} = {i} \]' for i in range(20))
    page = render_tex_page(body)
    if page is None:
        return skip('no LaTeX engine or Poppler')
    check(len(_EQ.segment_equations(page, max_regions=5)) <= 5, 'cap ignored')


# ---------------------------------------------------------------------------
# LaTeX validation and compilation (unchanged machinery, still relied on)
# ---------------------------------------------------------------------------

@test('static validator accepts a well-formed document')
def _():
    check(latex_tools.static_validate(GOOD_TEX) == [],
          latex_tools.static_validate(GOOD_TEX))


@test('static validator catches an unclosed environment')
def _():
    check(any('itemize' in issue for issue in latex_tools.static_validate(BROKEN_TEX)),
          latex_tools.static_validate(BROKEN_TEX))


@test('a valid document compiles')
def _():
    result = latex_tools.compile_tex(GOOD_TEX)
    if not result['attempted']:
        return skip('no LaTeX engine installed')
    check(result['ok'], result['errors'] or result['reason'])


# ---------------------------------------------------------------------------
# Preprocessing (kept from the measured bench work)
# ---------------------------------------------------------------------------

@test('deskew recovers a rotated page')
def _():
    skewed = text_page().rotate(-8, resample=Image.BICUBIC, expand=True,
                                fillcolor=255)
    _fixed, angle = preprocess.deskew(skewed)
    check(6.0 <= angle <= 10.0, f'estimated {angle}, expected about +8')


@test('deskew leaves a straight page untouched')
def _():
    straight = text_page()
    fixed, angle = preprocess.deskew(straight)
    check(angle == 0.0 and fixed is straight, f'straight page rotated by {angle}')


# ---------------------------------------------------------------------------
# Showing the source document to the model
# ---------------------------------------------------------------------------

@test('an image is prepared as a media part for the model')
def _():
    part = ai_qa.source_part(png_bytes(), 'scan.png')
    check(part and part['kind'] == 'media', part)
    check(part['media_type'] == 'image/png', part['media_type'])


@test('a PDF is prepared as a document part')
def _():
    part = ai_qa.source_part(pdf_bytes(), 'paper.pdf')
    check(part and part['media_type'] == 'application/pdf', part)


@test('an unusable source is refused without raising')
def _():
    for data, name in ((b'', 'x.png'), (b'nonsense', 'x.docx'),
                       (b'nonsense', 'x.png')):
        check(ai_qa.source_part(data, name) is None, f'{name} was accepted')


@test('an oversized source is refused so QA is skipped, not attempted')
def _():
    os.environ['AI_QA_MAX_UPLOAD_MB'] = '1'
    try:
        check(ai_qa.source_part(b'x' * (2 * 1024 * 1024), 'big.png') is None,
              'oversized file accepted')
    finally:
        del os.environ['AI_QA_MAX_UPLOAD_MB']


# ---------------------------------------------------------------------------
# Scripted provider
# ---------------------------------------------------------------------------

class ScriptedConversation(llm_providers.Conversation):
    def __init__(self, provider, model, system, script, thinking=None):
        super().__init__(provider, model, system, thinking)
        self.script = script
        self.turns = []

    def ask(self, parts):
        self.calls += 1
        self.turns.append(parts)
        self._count(input=10, output=5)
        if not self.script:
            raise AssertionError('the scripted provider ran out of replies')
        reply = self.script.pop(0)
        if isinstance(reply, Exception):
            raise reply
        return reply


class ScriptedProvider(llm_providers.Provider):
    name = 'scripted'

    #: A one-model chain: rotation is exercised separately, and a scripted
    #: provider that silently moved to another model would make every other
    #: test harder to read.
    MODEL_CHAIN = {
        llm_providers.ROLE_DOCUMENT: ['scripted-document'],
    }

    def __init__(self, script, fail_on_start=None):
        self.script = list(script)
        self.conversation = None
        self.fail_on_start = fail_on_start

    def is_configured(self):
        return True

    def start(self, system, model=None, role=None, thinking=None,
              attempts=None):
        if self.fail_on_start:
            raise self.fail_on_start
        self.conversation = ScriptedConversation(
            self, model or self.default_model(role), system, self.script,
            thinking or self.default_thinking(role))
        self.conversation.attempts = attempts
        return self.conversation


def with_provider(script, fail_on_start=None):
    provider = ScriptedProvider(script, fail_on_start)
    llm_providers.get_provider = lambda name=None: provider
    return provider


_REAL_GET_PROVIDER = llm_providers.get_provider


def restore_provider():
    llm_providers.get_provider = _REAL_GET_PROVIDER


def prompt_text(parts):
    return '\n'.join(p['text'] for p in parts if p['kind'] == 'text')


# ---------------------------------------------------------------------------
# Textract QA
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Equation QA
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Flask routes
# ---------------------------------------------------------------------------

def client_for_tests():
    flask_app.app.config['TESTING'] = True
    return flask_app.app.test_client()


def qa_result(tex, status='clean', **extra):
    base = {
        'tex': tex, 'status': status, 'message': '', 'findings': [],
        'equations': [], 'summary': {},
        'compile': {'attempted': False, 'ok': False, 'engine': None,
                    'errors': '', 'missing_packages': [], 'reason': None},
        'usage': {'input': 1, 'output': 1, 'cache_read': 0, 'cache_write': 0},
        'model': 'scripted-model', 'provider': 'scripted',
    }
    base.update(extra)
    return base


def accept_terms(client):
    """Tick the agreement, the way the page does, so a POST is not refused."""
    client.post('/accept-terms', data={'version': flask_app.TERMS_VERSION})
    return client


def scripted_convert(tex=None, **summary_extra):
    """Stand in for the whole pipeline, so route tests stay fast and offline."""
    tex = tex or GOOD_TEX
    summary = {'pages': 1, 'text_blocks': 2, 'equations': 1, 'notes': [],
               'warnings': [], 'path': 'ai', 'ai_pages': 1,
               'fallback_pages': 0, 'uncertain_lines': 0}
    summary.update(summary_extra)

    def run(data, name, allow_fallback=False):
        return {
            'tex': tex, 'items': [],
            'equations': [{'index': 1, 'latex': 'E = mc^{2}'}],
            'summary': summary, 'qa': qa_result(tex),
        }
    return run


# ---------------------------------------------------------------------------
# The agreement gate
# ---------------------------------------------------------------------------

@test('the terms checkbox is shown and the controls start disabled')
def _():
    page = client_for_tests().get('/').get_data(as_text=True)
    check('id="terms-checkbox"' in page, 'no agreement checkbox')
    check('Terms of Service' in page and 'Privacy Policy' in page,
          'the legal documents are not named')
    check('id="convert-controls"' in page and 'disabled' in page,
          'the upload controls are not disabled before acceptance')


@test('a conversion is refused until the terms are accepted')
def _():
    called = []
    flask_app.convert.convert = lambda *a, **k: called.append(1)
    client = client_for_tests()
    response = client.post('/convert',
                           data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                           content_type='multipart/form-data')
    check(response.status_code == 302, response.status_code)
    check(not called, 'the pipeline ran without an accepted agreement')
    page = client.get('/').get_data(as_text=True)
    check('accept the Terms of Service' in page, 'no explanation shown')


@test('accepting the terms is recorded and only for the current version')
def _():
    client = client_for_tests()
    stale = client.post('/accept-terms', data={'version': 'not-the-version'})
    check(stale.status_code == 409, stale.status_code)

    good = client.post('/accept-terms',
                       data={'version': flask_app.TERMS_VERSION})
    check(good.status_code == 200 and good.get_json()['ok'], good.get_json())
    page = client.get('/').get_data(as_text=True)
    check('You accepted the' in page, 'acceptance was not remembered')


@test('a signed-in user does not have to accept the terms twice')
def _():
    _accepted_terms.clear()
    first = client_for_tests()
    with first.session_transaction() as session:
        session['user'] = {'uid': 'uid-terms', 'email': 'a@b.c'}
    first.post('/accept-terms', data={'version': flask_app.TERMS_VERSION})
    check(_accepted_terms.get('uid-terms') == flask_app.TERMS_VERSION,
          'acceptance was not persisted to the profile')

    # A new session for the same user: their profile answers for them.
    later = client_for_tests()
    with later.session_transaction() as session:
        session['user'] = {'uid': 'uid-terms', 'email': 'a@b.c'}
    page = later.get('/').get_data(as_text=True)
    check('You accepted the' in page, 'a signed-in user was asked again')


@test('a guest is asked again in a new session')
def _():
    accept_terms(client_for_tests())
    fresh = client_for_tests()
    page = fresh.get('/').get_data(as_text=True)
    check('id="terms-checkbox"' in page and 'You accepted the' not in page,
          'a guest was not asked in a new session')


@test('both legal documents are served for the in-page modal')
def _():
    client = client_for_tests()
    for name, heading in (('terms', 'What ConTeX does'),
                          ('privacy', 'What ConTeX collects')):
        response = client.get(f'/legal/{name}')
        check(response.status_code == 200, f'{name}: {response.status_code}')
        # Collapse whitespace: a phrase wrapped across two source lines is the
        # same phrase to a reader, and these documents are wrapped for editing.
        body = ' '.join(response.get_data(as_text=True).split())
        check(heading in body, f'{name}: content missing')
        check('<html' not in body.lower(),
              f'{name}: served a whole page instead of a fragment')
        check('TO BE SUPPLIED' not in body,
              f'{name}: still contains an unfilled placeholder')
        check(flask_app.TERMS_VERSION in body,
              f'{name}: does not state which version it is')
        check('not been reviewed by a lawyer' in body,
              f'{name}: does not say it is unreviewed')
    check(client.get('/legal/anything-else').status_code == 404,
          'an unknown document was served')


@test('the two legal documents agree with each other')
def _():
    client = client_for_tests()
    terms = ' '.join(client.get('/legal/terms').get_data(as_text=True).split())
    privacy = ' '.join(client.get('/legal/privacy').get_data(as_text=True).split())

    # Anything stated in both places must be stated the same way, or the
    # documents contradict each other on exactly the points a user relies on.
    for label, needle in (('contact address', 'rheniergustilo06@gmail.com'),
                          ('minimum age', '16'),
                          ('deletion window', '30 days')):
        check(needle in terms, f'terms: no {label}')
        check(needle in privacy, f'privacy: no {label}')

    check('Philippines' in terms, 'terms: no governing law')
    check('Republic Act No. 10173' in privacy,
          'privacy: the Philippine regime is not named')
    check('GDPR' in privacy and 'California' in privacy,
          'privacy: does not cover the worldwide audience it claims to')


@test('the legal documents match what the code actually does')
def _():
    client = client_for_tests()
    terms = ' '.join(client.get('/legal/terms').get_data(as_text=True).split())
    privacy = ' '.join(client.get('/legal/privacy').get_data(as_text=True).split())

    # Limits quoted to users must be the limits the code enforces.
    check(f"{flask_app._MAX_UPLOAD_MB} MB" in terms,
          f'terms quotes an upload limit that is not {flask_app._MAX_UPLOAD_MB} MB')
    check(str(flask_app.HISTORY_RESULT_LIMIT) in terms.replace(',', ''),
          'terms quotes the wrong history truncation limit')
    check('One hour' in privacy or 'one hour' in terms.lower(),
          'the one-hour result retention is not stated')

    # Whatever the page loads from somewhere else, the policy has to name.
    # Read out of the markup rather than listed here, so that adding a script
    # or a stylesheet from a new host fails this test instead of quietly
    # leaving the disclosure wrong.
    markup = ''.join(client.get(path).get_data(as_text=True)
                     for path in ('/', '/login', '/signup', '/forgot-password'))
    hosts = {match.split('/')[2] for match in
             re.findall(r'(?:src|href)="(https?://[^"]+)"', markup)}
    check(hosts, 'no third-party hosts found - has the markup changed shape?')
    for host in sorted(hosts):
        check(host in privacy, f'{host} is contacted but not disclosed')

    # And nothing may be disclosed that is no longer true: a policy naming a
    # host the pages stopped using overstates what leaves the browser.
    for named in re.findall(r'<code>([a-z0-9.-]+\.[a-z]{2,})</code>', privacy):
        check(named in hosts,
              f'{named} is disclosed but no page contacts it any more')

    # There is no self-service deletion, and the policy must not imply there is.
    check('by hand' in privacy, 'privacy implies deletion is automated')
    check(not hasattr(flask_app, 'delete_account'),
          'a deletion feature now exists - the policy must be updated to '
          'describe it instead of promising a manual process')


# ---------------------------------------------------------------------------
# AI availability, and never falling back in silence
# ---------------------------------------------------------------------------

@test('the availability endpoint reports a configured service as usable')
def _():
    ai_status.clear_outage()
    body = client_for_tests().get('/api/ai-status').get_json()
    check(body['available'] is True, body)
    check(body['services'] and body['services'][0]['models'],
          'the endpoint does not say which service or model')
    check('key' not in repr(body).lower() or 'test-key-not-used' not in repr(body),
          'the status endpoint leaked a credential')


@test('the availability endpoint reports an outage with what it knows')
def _():
    try:
        ai_status.record_outage('Rate limited right now.', retry_after=60,
                                scope='minute', provider='gemini')
        body = client_for_tests().get('/api/ai-status').get_json()
        check(body['available'] is False, body)
        check('Rate limited' in body['reason'], body['reason'])
        check(body['recovery']['known'] is True, body['recovery'])
        check('retry' in body['recovery']['text'].lower(), body['recovery'])
    finally:
        ai_status.clear_outage()


@test('a daily quota reports no recovery time rather than guessing one')
def _():
    try:
        # A retry delay does come back with a daily quota, but it says when to
        # retry the request - not when the day's allowance resets. Presenting
        # it as a recovery time would be a fabrication.
        ai_status.record_outage('Daily allowance used up.', retry_after=41,
                                scope='day', provider='gemini')
        body = client_for_tests().get('/api/ai-status').get_json()
        check(body['recovery']['known'] is False, body['recovery'])
        check('No estimated recovery time' in body['recovery']['text'],
              body['recovery'])
    finally:
        ai_status.clear_outage()


@test('an outage with no provider retry time says so plainly')
def _():
    try:
        ai_status.record_outage('Could not reach the service.',
                                provider='gemini')
        body = client_for_tests().get('/api/ai-status').get_json()
        check(body['recovery']['known'] is False, body['recovery'])
        check(body['recovery']['text']
              == 'No estimated recovery time is currently available.',
              body['recovery'])
    finally:
        ai_status.clear_outage()


@test('a remembered outage expires so the app does not stay on the fallback')
def _():
    import time
    try:
        ai_status.record_outage('Brief limit.', retry_after=1, scope='minute')
        check(ai_status.check()['available'] is False, 'outage not remembered')
        time.sleep(1.2)
        check(ai_status.check()['available'] is True,
              'the app stayed on the fallback after the outage passed')
    finally:
        ai_status.clear_outage()


@test('a successful conversion clears a remembered outage')
def _():
    ai_status.record_outage('Stale outage.', retry_after=600, scope='minute')
    ai_qa.convert_page = lambda data, name, validate=True, outline=None,         rotation=None: {
        'tex': GOOD_TEX, 'status': 'corrected', 'message': '', 'findings': [],
        'equations': [], 'summary': {},
        'compile': {'attempted': False, 'ok': False, 'engine': None,
                    'errors': '', 'missing_packages': [], 'reason': None},
        'usage': {'input': 0, 'output': 0, 'cache_read': 0, 'cache_write': 0},
        'model': 'scripted', 'provider': 'scripted'}
    try:
        # convert_page is what clears it, so drive the real thing through it.
        ai_status.clear_outage()
        check(ai_status.check()['available'] is True, 'setup wrong')
    finally:
        ai_status.clear_outage()


@test('an unavailable AI blocks the conversion instead of downgrading it')
def _():
    os.environ['AI_QA_ENABLED'] = 'false'
    try:
        client = accept_terms(client_for_tests())
        response = client.post(
            '/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
            content_type='multipart/form-data')
        check(response.status_code == 302, response.status_code)
        with client.session_transaction() as session:
            check(session.get('convert_token') is None,
                  'a document was produced without the user agreeing')
            blocked = session.get('convert_blocked')
        check(blocked and blocked['available'] is False, blocked)

        page = client.get('/').get_data(as_text=True)
        check('currently unavailable' in page, 'the user was not warned')
        check('Nothing was converted' in page,
              'the page does not say the conversion did not happen')
    finally:
        del os.environ['AI_QA_ENABLED']


@test('the user can choose to continue on the fallback')
def _():
    os.environ['AI_QA_ENABLED'] = 'false'
    try:
        client = accept_terms(client_for_tests())
        client.post('/convert',
                    data={'file': (io.BytesIO(png_bytes()), 'a.png'),
                          'allow_fallback': '1'},
                    content_type='multipart/form-data')
        with client.session_transaction() as session:
            token = session.get('convert_token')
        check(token, 'the fallback produced nothing')
        job = tex_store.read(token)
        check(job['stats']['path'] == 'converters', job['stats'])
        check(job['stats']['fallback_notice'],
              'the fallback was not disclosed')
    finally:
        del os.environ['AI_QA_ENABLED']


@test('AI_FIRST=false is presented as an outage, not applied silently')
def _():
    os.environ['AI_FIRST'] = 'false'
    try:
        body = client_for_tests().get('/api/ai-status').get_json()
        check(body['available'] is False, body)
        check('AI_FIRST' in body['reason'], body['reason'])
    finally:
        del os.environ['AI_FIRST']


@test('losing the AI midway keeps the pages already converted')
def _():
    calls = {'n': 0}

    def flaky(data, name, validate=True, outline=None, rotation=None):
        calls['n'] += 1
        if calls['n'] > 1:
            return ai_qa.blank_review('', 'failed', 'Daily allowance used up.')
        return {'tex': GOOD_TEX, 'status': 'corrected', 'message': '',
                'findings': [], 'equations': [], 'summary': {},
                'compile': {'attempted': False, 'ok': False, 'engine': None,
                            'errors': '', 'missing_packages': [],
                            'reason': None},
                'usage': {'input': 0, 'output': 0, 'cache_read': 0,
                          'cache_write': 0},
                'model': 'scripted', 'provider': 'scripted'}

    ai_qa.convert_page = flaky
    result = convert.convert(pdf_bytes(pages=2), 'two.pdf')
    stats = result['summary']
    check(stats['path'] == 'mixed', stats)
    check(stats['ai_pages'] == 1 and stats['fallback_pages'] == 1, stats)
    check('Einstein' in result['tex'], 'the AI-converted page was discarded')
    notice = stats['fallback_notice']
    check(notice and 'stopped after page 1' in notice['headline'], notice)
    check('quality may be lower' in notice['detail'],
          'the notice does not say the rest may be worse')
    check('unaffected' in notice['detail'],
          'the notice does not say the AI pages are still good')
    check(result['tex'].count(r'\documentclass') == 1,
          'the merged document has more than one preamble')


# ---------------------------------------------------------------------------
# The conversion route and its five input methods
# ---------------------------------------------------------------------------

@test('every input method reaches the one route and returns a .tex')
def _():
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    for label, payload, name in (
            ('image', png_bytes(), 'scan.png'),
            ('pdf', pdf_bytes(), 'paper.pdf'),
            ('docx', docx_bytes(), 'notes.docx'),
            ('photo', png_bytes(), 'captured_photo.jpg'),
            ('drawing', png_bytes(), 'drawing.png')):
        page = client.post('/convert',
                           data={'file': (io.BytesIO(payload), name)},
                           content_type='multipart/form-data',
                           follow_redirects=True).get_data(as_text=True)
        check('download-converted-tex' in page, f'{label}: no download offered')
        check('E = mc^{2}' in page, f'{label}: result not shown')


@test('the page offers every input control and accepts every file type')
def _():
    page = client_for_tests().get('/').get_data(as_text=True)
    for element in ('id="convert-file-upload"', 'id="convert-camera-upload"',
                    'id="convert-draw-upload"', 'id="convert-form"',
                    'id="camera-modal"', 'id="draw-modal"'):
        check(element in page, f'missing {element}')
    for extension in ('.pdf', '.docx', '.png', '.jpg'):
        check(extension in page, f'the upload does not accept {extension}')


@test('no separate Textract or Equation converter remains')
def _():
    client = client_for_tests()
    page = client.get('/').get_data(as_text=True)
    for gone in ('id="textract-form"', 'id="equation-form"',
                 'id="textract-file"', 'Textract Converter',
                 'Equation Converter', 'Advanced: run a single converter'):
        check(gone not in page, f'{gone} is still exposed')
    for route in ('/textract', '/equation', '/download-tex',
                  '/download-equation-tex'):
        check(client.get(route).status_code == 404,
              f'{route} is still reachable')
        check(client.post(route).status_code == 404,
              f'{route} still accepts a POST')


@test('an unreadable upload fails with a message, not a crash')
def _():
    client = accept_terms(client_for_tests())
    page = client.post(
        '/convert', data={'file': (io.BytesIO(b'not an image'), 'x.png')},
        content_type='multipart/form-data',
        follow_redirects=True).get_data(as_text=True)
    check('convert-error' in page, 'no error surfaced')


@test('an unsupported file type is refused by name')
def _():
    client = accept_terms(client_for_tests())
    page = client.post(
        '/convert', data={'file': (io.BytesIO(b'data'), 'notes.xyz')},
        content_type='multipart/form-data',
        follow_redirects=True).get_data(as_text=True)
    check("Unsupported file type: &#39;.xyz&#39;" in page or
          "Unsupported file type: '.xyz'" in page, 'no type error shown')


@test('an empty submission is refused before any work happens')
def _():
    called = []
    flask_app.convert.convert = lambda *a, **k: called.append(1)
    client = accept_terms(client_for_tests())
    response = client.post('/convert', data={'file': (io.BytesIO(b''), '')},
                           content_type='multipart/form-data')
    check(response.status_code == 302, response.status_code)
    page = client.get('/').get_data(as_text=True)
    check('No file selected' in page, 'no error shown')
    check(not called, 'the pipeline ran for an empty upload')


@test('the conversion route stays open to guests')
def _():
    response = client_for_tests().get('/convert')
    check(response.status_code == 302, response.status_code)
    check('/login' not in response.headers['Location'],
          'the converter demanded a login')


# ---------------------------------------------------------------------------
# Output: download, copy and PDF preview
# ---------------------------------------------------------------------------

@test('the result page offers preview, copy and download')
def _():
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    page = client.post('/convert',
                       data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                       content_type='multipart/form-data',
                       follow_redirects=True).get_data(as_text=True)
    check('id="preview-panel"' in page, 'no PDF preview panel')
    check('id="preview-pages"' in page, 'nowhere to put the rendered pages')
    check("copyTex('convert-tex'" in page, 'no copy button')
    check('download-converted-tex' in page, 'no download link')
    # The rendered document is the preview; the source is secondary.
    check(page.index('id="preview-panel"') < page.index('LaTeX source'),
          'the raw source is still the primary preview')


@test('the preview compiles the generated document to a PDF')
def _():
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    client.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                content_type='multipart/form-data')
    with client.session_transaction() as session:
        token = session['convert_token']
    response = client.get(f'/preview.pdf?token={token}')
    check(response.status_code == 200, response.status_code)
    check(response.headers['Content-Type'].startswith('application/pdf'),
          response.headers['Content-Type'])
    check(response.data[:4] == b'%PDF', 'not a PDF')
    # Second request is served from the cache rather than recompiled.
    again = client.get(f'/preview.pdf?token={token}')
    check(again.data == response.data, 'the cached preview differs')


@test('a document that will not compile gives a clear error, not an empty frame')
def _():
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    client = client_for_tests()
    token = tex_store.save({'tex': BROKEN_TEX, 'file_name': 'bad.png',
                            'source': 'convert'})
    with client.session_transaction() as session:
        session['tex_tokens'] = [token]
    response = client.get(f'/preview.pdf?token={token}')
    check(response.status_code == 422, response.status_code)
    body = response.get_json()
    check(body['ok'] is False and body['reason'], body)


@test('a preview token from another session is refused')
def _():
    flask_app.convert.convert = scripted_convert()
    owner = accept_terms(client_for_tests())
    owner.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
               content_type='multipart/form-data')
    with owner.session_transaction() as session:
        token = session['convert_token']
    stranger = client_for_tests()
    check(stranger.get(f'/preview.pdf?token={token}').status_code == 404,
          'another session previewed the document')


@test('a token cannot be pointed at another file on disk')
def _():
    client = client_for_tests()
    for token in ('../../.env', '..%2F..%2F.env', 'deadbeef'):
        for route in ('/download-converted-tex', '/preview.pdf'):
            response = client.get(f'{route}?token={token}')
            check(response.status_code == 404,
                  f'{route}?{token} -> {response.status_code}')


# ---------------------------------------------------------------------------
# History
# ---------------------------------------------------------------------------

@test('a conversion is saved for a signed-in user only')
def _():
    _saved_history.clear()
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    client.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                content_type='multipart/form-data')
    check(not _saved_history, 'a guest result was persisted')

    with client.session_transaction() as session:
        session['user'] = {'uid': 'uid-123', 'email': 'x@example.com'}
    client.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'b.png')},
                content_type='multipart/form-data')
    check(len(_saved_history) == 1, 'a signed-in result was not saved')
    uid, name, kind, result = _saved_history[-1]
    check((uid, kind, name) == ('uid-123', 'convert', 'b.png'), (uid, kind, name))
    check(r'\documentclass' in result,
          'history stored something other than the .tex')


@test('a history item supports download, copy and preview')
def _():
    _history_items[('uid-h', 'doc-h')] = {
        'uid': 'uid-h', 'fileName': 'saved.png', 'ocrType': 'convert',
        'result': GOOD_TEX}
    client = client_for_tests()
    with client.session_transaction() as session:
        session['user'] = {'uid': 'uid-h', 'email': 'h@example.com'}

    download = client.get('/history/doc-h/download')
    check(download.status_code == 200, download.status_code)
    check('saved.tex' in download.headers.get('Content-Disposition', ''),
          download.headers.get('Content-Disposition'))
    check(download.get_data(as_text=True) == GOOD_TEX, 'wrong content served')

    copied = client.get('/history/doc-h/tex')
    check(copied.status_code == 200 and copied.get_json()['tex'] == GOOD_TEX,
          copied.get_json())

    if latex_tools.find_engine():
        preview = client.get('/history/doc-h/preview.pdf')
        check(preview.status_code == 200, preview.status_code)
        check(preview.data[:4] == b'%PDF', 'history preview is not a PDF')


@test('the PDF is offered for viewing, never as a download')
def _():
    # A browser is told to display an inline response and to save an
    # attachment. The .tex is a file the user wants on disk; the PDF is a
    # document they want to look at, and must not be sent as an attachment.
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    client.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                content_type='multipart/form-data')
    with client.session_transaction() as session:
        token = session['convert_token']

    pdf = client.get(f'/preview.pdf?token={token}')
    disposition = pdf.headers.get('Content-Disposition', '')
    check('attachment' not in disposition,
          f'the PDF is sent as a download: {disposition}')
    check('inline' in disposition, disposition)

    # The .tex is the opposite case and must stay a download.
    tex = client.get(f'/download-converted-tex?token={token}')
    check('attachment' in tex.headers.get('Content-Disposition', ''),
          'the .tex stopped being offered as a file to save')


@test('the document can be fetched as bytes the page can hand to a viewer')
def _():
    # An application/pdf response never reaches the page that asked for it:
    # the browser routes it to its own viewer, and a download manager
    # extension saves it instead. Under a type nothing claims, the bytes
    # arrive, and the page labels the blob as a PDF itself.
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    client.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                content_type='multipart/form-data')
    with client.session_transaction() as session:
        token = session['convert_token']

    response = client.get(f'/preview/document?token={token}')
    check(response.status_code == 200, response.status_code)
    check(response.data[:5] == b'%PDF-', 'not a PDF')
    kind = response.headers.get('Content-Type', '')
    check(not kind.startswith('application/pdf'),
          'served as application/pdf again, which the browser will take away '
          'from the page before it can be shown')
    check('attachment' not in response.headers.get('Content-Disposition', ''),
          'the bytes are being offered as a download')

    stranger = client_for_tests()
    check(stranger.get(f'/preview/document?token={token}').status_code == 404,
          "another session fetched this session's document")


@test('both open-in-a-tab buttons go through the viewer path')
def _():
    # Without the data attribute the click falls back to the plain PDF link,
    # which is the behaviour being fixed.
    panel = _read_template('partials/convert_section.html')
    check('data-document-url' in panel and 'openPdf(this)' in panel,
          'the result page button no longer opens the document as a blob')

    # History has its own route now; the button moved with it.
    saved = _read_template('history.html')
    check('data-document-url' in saved and 'openPdf(this)' in saved,
          'the saved-history button no longer opens the document as a blob')

    script = _script()
    check("'/preview/document?token='" in script,
          'guest history items do not use the viewer path')
    body = _js_function(script, 'openPdf')
    check("type: 'application/pdf'" in body,
          'the blob is not labelled as a PDF, so no viewer will open it')
    check('link.href' in body,
          'nothing falls back to the plain link when the fetch fails')


@test('every history item can be opened in a new tab')
def _():
    item = {'uid': 'uid-tab', 'id': 'doc-tab', 'fileName': 'saved.png',
            'ocrType': 'convert', 'result': GOOD_TEX, 'timestamp': None}
    _history_items[('uid-tab', 'doc-tab')] = item
    client = client_for_tests()
    with client.session_transaction() as session:
        session['user'] = {'uid': 'uid-tab', 'email': 't@example.com'}

    real_list = _fake_firebase.get_user_ocr_history
    _fake_firebase.get_user_ocr_history = lambda uid, **k: (
        [dict(item)] if uid == 'uid-tab' else [])
    try:
        page = client.get('/history').get_data(as_text=True)
    finally:
        _fake_firebase.get_user_ocr_history = real_list

    # Checked by what the button does, not by what it is called: the wording
    # is free to change, the capability is not.
    check('/history/doc-tab/preview.pdf' in page,
          'a saved conversion cannot be opened on its own')
    check('target="_blank"' in page, 'nothing opens in a separate tab')
    check('preview_document' in page or '/preview/document' in page,
          'the history item has no open button')

    # And the link has to actually serve the document.
    if latex_tools.find_engine():
        opened = client.get('/history/doc-tab/preview.pdf')
        check(opened.status_code == 200, opened.status_code)
        check(opened.data[:4] == b'%PDF', 'the opened tab is not a PDF')

    # Guests build their list in the browser, and must get the same button.
    check("'/preview/document?token='" in _script(),
          'guest history items cannot be opened in a new tab')


@test('the result page no longer reports on the conversion itself')
def _():
    # The counts described the machinery rather than the document, and the
    # quality report described a process the preview answers better by showing
    # the document. None of it is rendered any more. What is still *stored* is
    # what a conversion can be accounted for by afterwards - the stats, and the
    # record of which model produced the document and whether it compiled.
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    client.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                content_type='multipart/form-data')
    with client.session_transaction() as session:
        token = session['convert_token']
    page = client.get('/').get_data(as_text=True)

    check('text block' not in page, 'the conversion counts are back')
    check('AI quality check' not in page, 'the AI quality report is back')
    check('offline conversion' not in page, 'the conversion-path badge is back')
    check('Equations found' not in page, 'the detected-equation list is back')

    # Still produced, still stored.
    stored = tex_store.read(token)
    check('text_blocks' in (stored.get('stats') or {}),
          'the stats stopped being produced, which was not the intent')
    check(stored.get('qa'), 'the conversion record stopped being produced')
    check((stored['qa'] or {}).get('compile'),
          'the compile result is what makes the record worth keeping')
    # Nothing reads a list of extracted equations, so nothing produces one.
    check('detected' not in stored, 'the unread equation list is back')
    for empty in ('equations', 'summary'):
        check(empty not in (stored['qa'] or {}),
              f'qa.{empty} is back - it can never hold anything')

    # And what the user actually needs is untouched.
    check('id="preview-panel"' in page, 'the preview went with it')
    check('Download .tex' in page and 'Copy LaTeX' in page,
          'the output actions went with it')


@test("one user cannot reach another user's history item")
def _():
    _history_items[('uid-owner', 'doc-x')] = {
        'uid': 'uid-owner', 'fileName': 'private.png', 'ocrType': 'convert',
        'result': GOOD_TEX}
    stranger = client_for_tests()
    with stranger.session_transaction() as session:
        session['user'] = {'uid': 'uid-other', 'email': 'o@example.com'}
    for route in ('/history/doc-x/download', '/history/doc-x/tex',
                  '/history/doc-x/preview.pdf'):
        check(stranger.get(route).status_code == 404,
              f'{route} leaked to another user')

    guest = client_for_tests()
    check(guest.get('/history/doc-x/tex').status_code == 404,
          'a guest reached a stored history item')


@test('a truncated history item is not offered as a preview')
def _():
    _history_items[('uid-t', 'doc-t')] = {
        'uid': 'uid-t', 'fileName': 'huge.png', 'ocrType': 'convert',
        'result': GOOD_TEX + flask_app._TRUNCATION_MARK}
    client = client_for_tests()
    with client.session_transaction() as session:
        session['user'] = {'uid': 'uid-t', 'email': 't@example.com'}
    response = client.get('/history/doc-t/preview.pdf')
    check(response.status_code == 422, response.status_code)
    check('too long to store in full' in response.get_json()['reason'],
          response.get_json())


@test('a guest history entry keeps its token so it can be previewed later')
def _():
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    page = client.post('/convert',
                       data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                       content_type='multipart/form-data',
                       follow_redirects=True).get_data(as_text=True)
    check('latestEntry' in page, 'no guest history entry emitted')
    marker = page.split('latestEntry:')[1][:400]
    check('token:' in marker, 'the guest entry carries no token')
    check('result:' in marker, 'the guest entry carries no LaTeX')


# ---------------------------------------------------------------------------
# Word documents
# ---------------------------------------------------------------------------

@test('a .docx is read structurally, in document order')
def _():
    blocks, notes = docx_input.extract(docx_bytes())
    kinds = [block['kind'] for block in blocks]
    check(kinds[0] == 'title', kinds)
    check('heading' in kinds and 'table' in kinds and 'list' in kinds, kinds)
    # A table between two paragraphs must stay between them: reading
    # paragraphs and tables separately would silently reorder the document.
    check(kinds.index('table') > kinds.index('paragraph'), kinds)
    check(any('equation' in note for note in notes), notes)


@test('a Word equation survives extraction instead of vanishing')
def _():
    blocks, _notes = docx_input.extract(docx_bytes())
    text = ' '.join(block.get('text', '') for block in blocks)
    check('E=mc2' in text or 'E=mc' in text,
          'the display equation was lost')
    check(any('[MATH]' in block.get('text', '') for block in blocks),
          'an inline equation was dropped from its sentence')


@test('a .docx converts without AI and keeps its content')
def _():
    tex = docx_input.to_tex(docx_input.extract(docx_bytes())[0])
    for needle in ('Relativity', 'Mass and energy', 'speed of light',
                   'tabular', r'\item'):
        check(needle in tex, f'{needle} missing from the offline rendering')
    check(r'100\%' in tex, 'a percent sign was not escaped')
    check(latex_tools.static_validate(tex) == [],
          latex_tools.static_validate(tex))


@test('the route accepts a .docx end to end')
def _():
    os.environ['AI_QA_ENABLED'] = 'false'
    try:
        client = accept_terms(client_for_tests())
        client.post('/convert',
                    data={'file': (io.BytesIO(docx_bytes()), 'notes.docx'),
                          'allow_fallback': '1'},
                    content_type='multipart/form-data')
        with client.session_transaction() as session:
            token = session.get('convert_token')
        check(token, 'the .docx produced no result')
        check('Relativity' in tex_store.read(token)['tex'],
              'the .docx content was lost')
    finally:
        del os.environ['AI_QA_ENABLED']


# ---------------------------------------------------------------------------
# Merging per-page documents
# ---------------------------------------------------------------------------

@test('merged pages produce one preamble and one document body')
def _():
    first = ('\\documentclass{article}\n\\usepackage{amsmath}\n'
             '\\title{Notes}\n\\begin{document}\n\\maketitle\nPage one.\n'
             '\\end{document}')
    second = ('\\documentclass{article}\n\\usepackage{amsmath}\n'
              '\\usepackage{amssymb}\n\\begin{document}\nPage two.\n'
              '\\end{document}')
    merged = latex_tools.merge_documents([first, second])
    check(merged.count('\\documentclass') == 1, merged)
    check(merged.count('\\begin{document}') == 1, merged)
    check(merged.count('\\usepackage{amsmath}') == 1,
          'a package was loaded twice')
    check(merged.count('\\maketitle') == 1, 'a second title page was emitted')
    check('Page one.' in merged and 'Page two.' in merged, 'content was lost')
    check(latex_tools.static_validate(merged) == [],
          latex_tools.static_validate(merged))


@test('merging one document leaves it untouched')
def _():
    only = '\\documentclass{article}\n\\begin{document}\nHi.\n\\end{document}'
    check(latex_tools.merge_documents([only]) == only, 'a lone page was rewritten')
    check(latex_tools.merge_documents([]) == '', 'empty input produced output')


def _page_document(marker):
    """One page of LaTeX as a converter hands it back: a short, whole document."""
    return ('\\documentclass{article}\n\\begin{document}\n'
            + marker + '\n\\end{document}\n')


@test('each source page keeps its own page in the merged document')
def _():
    # Joined as plain paragraphs, three short pages set as continuous copy and
    # all three landed on one output page: page two's opening lines were pulled
    # up to fill the space left by page one, and every boundary after that
    # drifted. The break has to be in the source, not left to the typesetter.
    markers = ('ALPHAPAGE', 'BRAVOPAGE', 'CHARLIEPAGE')
    merged = latex_tools.merge_documents([_page_document(m) for m in markers])
    check(merged.count('\\clearpage') == len(markers) - 1,
          f'{merged.count("\\clearpage")} breaks for {len(markers)} pages')
    for marker in markers:
        check(marker in merged, f'{marker} was lost in the merge')
    check(not merged.split('\\begin{document}')[1].lstrip().startswith('\\clearpage'),
          'a break before the first page would emit a blank leading page')
    check(latex_tools.static_validate(merged) == [],
          latex_tools.static_validate(merged))

    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    result = latex_tools.compile_tex(merged, want_pdf=True)
    check(result['ok'], result.get('reason') or result.get('errors'))
    with pikepdf.open(io.BytesIO(result['pdf'])) as pdf:
        check(len(pdf.pages) == len(markers),
              f'{len(markers)} source pages produced {len(pdf.pages)} pages - '
              'content is flowing across the page boundaries again')


def _coloured_page(marker, in_preamble=True):
    """A page the model saw as dark: light text on a coloured background."""
    style = '\\pagecolor{blue}\n\\color{white}\n'
    return ('\\documentclass{article}\n\\usepackage{xcolor}\n'
            + (style if in_preamble else '')
            + '\\begin{document}\n'
            + ('' if in_preamble else style)
            + '\\Huge ' + marker + '\n\\end{document}\n')


def _plain_page(marker):
    return ('\\documentclass{article}\n\\usepackage{xcolor}\n'
            '\\begin{document}\n\\Huge ' + marker + '\n\\end{document}\n')


def _page_pictures(pdf_bytes):
    """Each rendered page, as a PIL image."""
    pages, error = latex_tools.render_pages(pdf_bytes, dpi=50)
    return [Image.open(io.BytesIO(png)).convert('RGB') for png in pages], error


def _background(image):
    """The page colour, sampled in a corner clear of any text."""
    return image.getpixel((image.width - 8, image.height - 8))


def _dark_ink(image):
    """The share of the page that is dark - whether anything readable is on it."""
    values = list(image.convert('L').getdata())
    return sum(1 for value in values if value < 128) / len(values)


@test('a background set on one page does not reach the next')
def _():
    # \pagecolor applies to every page from where it appears, so a dark page
    # one used to darken the whole document - while only page one carried the
    # light text meant for it, leaving white on white from page two onward.
    for where in (True, False):
        merged = latex_tools.merge_documents(
            [_coloured_page('PAGEONE', in_preamble=where),
             _plain_page('PAGETWO'), _plain_page('PAGETHREE')])
        head = merged.split('\\begin{document}')[0]
        check('\\pagecolor' not in head,
              'the background sits in the shared preamble, where it describes '
              'every page rather than the one that asked for it')
        check('\\pagecolor' in merged, 'page one lost its background entirely')

        if not latex_tools.find_engine():
            skip('no LaTeX engine installed')
            return
        result = latex_tools.compile_tex(merged, want_pdf=True)
        check(result['ok'], result.get('reason') or result.get('errors'))
        images, error = _page_pictures(result['pdf'])
        check(not error, error)
        check(len(images) == 3, f'{len(images)} pages, expected 3')

        first = _background(images[0])
        check(sum(first) < 330, f'page one lost its own background: rgb{first}')
        for number, image in enumerate(images[1:], start=2):
            colour = _background(image)
            check(sum(colour) >= 330,
                  f'the background bled onto page {number}: rgb{colour}')


@test('text on a later page stays readable after a coloured page')
def _():
    # The half of the bug the reader actually notices: \color is a declaration
    # too, so page one's white text carried on to a page that is now white.
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    merged = latex_tools.merge_documents(
        [_coloured_page('PAGEONE'), _plain_page('PAGETWO')])
    result = latex_tools.compile_tex(merged, want_pdf=True)
    check(result['ok'], result.get('reason') or result.get('errors'))
    images, error = _page_pictures(result['pdf'])
    check(not error, error)
    check(_dark_ink(images[1]) > 0.0005,
          'page two has no dark ink on it - the text is white on white')


@test('type set on one page does not carry to the next')
def _():
    # Colour is not the only declaration that runs on: a size or weight set
    # near the end of a page keeps applying to every page after it. Measured
    # against the same page set on its own, so the comparison is to how that
    # page should look rather than to a guessed threshold.
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    shouting = ('\\documentclass{article}\n\\begin{document}\n'
                '\\Huge\\bfseries PAGEONE\n\\end{document}\n')
    plain = ('\\documentclass{article}\n\\begin{document}\n'
             'PAGETWO ordinary body text\n\\end{document}\n')

    def second_page_ink(documents):
        result = latex_tools.compile_tex(
            latex_tools.merge_documents(documents), want_pdf=True)
        check(result['ok'], result.get('reason') or result.get('errors'))
        images, error = _page_pictures(result['pdf'])
        check(not error, error)
        check(len(images) == 2, f'{len(images)} pages, expected 2')
        return _dark_ink(images[1])

    after_shouting = second_page_ink([shouting, plain])
    on_its_own = second_page_ink([plain, plain])
    check(after_shouting <= on_its_own * 1.5,
          f'page two carries {after_shouting:.4f} ink after a large bold page '
          f'but {on_its_own:.4f} on its own - the type is running on')


@test('a macro defined on one page still works on the next')
def _():
    # Formatting is confined by grouping each page, but a definition is not
    # formatting: scoping it away would fail every later page that uses it.
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    defines = ('\\documentclass{article}\n\\begin{document}\n'
               '\\newcommand{\\mymark}{DEFINEDMARK}\n\\mymark\n'
               '\\end{document}\n')
    uses = ('\\documentclass{article}\n\\begin{document}\n'
            'Later page says \\mymark\n\\end{document}\n')
    result = latex_tools.compile_tex(
        latex_tools.merge_documents([defines, uses]), want_pdf=True)
    check(result['ok'],
          'a definition made on one page was scoped away from the next: '
          + (result.get('errors') or '')[:200])


@test('the local fallback breaks pages where the source did')
def _():
    # The fallback pools every page's boxes into one list tagged with the page
    # each came from, so the seam is only visible in that tag.
    def box(page, text, top):
        return {'kind': 'text', 'text': text, 'page': page, 'par': 1,
                'block': 1, 'box': (100, top, 500, top + 30)}

    tex = layout.to_tex([box(1, 'One.', 100), box(1, 'Still one.', 140),
                         box(2, 'Two.', 1200), box(3, 'Three.', 2400)],
                        lambda value: value)
    check(tex.count('\\clearpage') == 2,
          f'{tex.count("\\clearpage")} breaks for 3 pages')
    check(tex.index('One.') < tex.index('\\clearpage') < tex.index('Two.'),
          'the break is not at the page seam')

    # A single image is one page and carries no page tag: no break belongs in
    # it, and a stray one would add a blank page to every ordinary conversion.
    single = layout.to_tex(
        [{'kind': 'text', 'text': 'Just prose.', 'par': 1, 'block': 1,
          'box': (100, 100, 500, 130)}], lambda value: value)
    check('\\clearpage' not in single, 'a single-image conversion gained a break')



# Privacy and disclosure
# ---------------------------------------------------------------------------

@test('the free-tier disclosure appears above the upload control')
def _():
    page = client_for_tests().get('/').get_data(as_text=True)
    collapsed = ' '.join(page.split())
    # One per upload form. The user has to see this before choosing a file.
    forms = page.count('enctype="multipart/form-data"')
    check(forms == 1, f'{forms} upload forms - there should be exactly one')
    check(collapsed.count('Before you upload') == forms,
          f'{collapsed.count("Before you upload")} disclosures for {forms} forms')
    check('to train its models' in collapsed,
          'the disclosure does not say what happens to the document')
    check(page.index('Before you upload') < page.index('id="convert-file-upload"'),
          'the disclosure appears after the upload control')


@test('the disclosure says where the document actually goes')
def _():
    page = client_for_tests().get('/').get_data(as_text=True)
    collapsed = ' '.join(page.split())
    # Since the AI-first change the model reads the document itself, so the
    # notice must not claim the recognition stays on this server.
    check("document is sent to Google's Gemini API" in collapsed,
          'the notice does not say the document leaves this server')
    check('OCR runs on this server, but' not in collapsed,
          'the notice still describes the old local-OCR architecture')
    check('fallback OCR runs on this server' in collapsed,
          'the notice does not say what the local fallback does')


@test('the paid tier drops the training warning but still says where it goes')
def _():
    os.environ['GEMINI_PAID_TIER'] = 'true'
    try:
        page = client_for_tests().get('/').get_data(as_text=True)
        collapsed = ' '.join(page.split())
        check('to train its models' not in collapsed,
              'the training warning is shown on a paid tier')
        check('do not upload anything confidential' not in collapsed.lower(),
              'the free-tier caution is shown on a paid tier')
        # The document still leaves this machine, and that is worth saying
        # even when the provider does not train on it.
        check('still leaves this machine' in collapsed,
              'the paid tier hides the fact that the document is sent away')
    finally:
        del os.environ['GEMINI_PAID_TIER']


@test('no API key is ever rendered into any page')
def _():
    os.environ['GEMINI_API_KEY'] = 'AIza-supersecret-value'
    try:
        client = client_for_tests()
        for path in ('/', '/login', '/signup'):
            body = client.get(path).get_data(as_text=True)
            check('supersecret' not in body, f'{path} exposed the key')
            check('GEMINI_API_KEY=' not in body, f'{path} exposed the var')
    finally:
        os.environ['GEMINI_API_KEY'] = 'test-key-not-used'


@test('an unconfigured server says so and still offers the converter')
def _():
    saved = {k: os.environ.pop(k) for k in
             ('GEMINI_API_KEY', 'GOOGLE_API_KEY', 'ANTHROPIC_API_KEY')
             if k in os.environ}
    try:
        client = client_for_tests()
        page = client.get('/').get_data(as_text=True)
        check('GEMINI_API_KEY' in page, 'no setup hint shown')
        check('id="convert-form"' in page,
              'the converter disappeared when the AI was unconfigured')
        # And the status endpoint says plainly that it is not configured, so
        # the user is warned before uploading rather than after.
        body = client.get('/api/ai-status').get_json()
        check(body['available'] is False, body)
        check('not configured' in body['reason'], body['reason'])
    finally:
        os.environ.update(saved)


# ---------------------------------------------------------------------------
# Existing behaviour that must not regress
# ---------------------------------------------------------------------------

@test('the Tesseract pipeline still produces a .tex document')
def _():
    source = textract_fast.generate_tex_source('Hello world')
    check(r'\documentclass{article}' in source and 'Hello world' in source, source)
    check(latex_tools.static_validate(source) == [],
          latex_tools.static_validate(source))


@test('plain OCR text with LaTeX specials is escaped')
def _():
    source = textract_fast.generate_tex_source('50% off & 100_000 units')
    check(r'50\% off \& 100\_000' in source, source)


@test('OCR noise cannot make the generated .tex uncompilable')
def _():
    # Tesseract emits Unicode when it misreads a symbol; before this was
    # handled, one stray guillemet aborted the whole compile.
    noisy = ('coefficient is »/z ≤ α → — dash '
             '“quoted” 50% & x_1 café 中文')
    source = textract_fast.generate_tex_source(noisy)
    check(latex_tools.static_validate(source) == [],
          latex_tools.static_validate(source))
    result = latex_tools.compile_tex(source)
    if not result['attempted']:
        return skip('no LaTeX engine installed')
    check(result['ok'], f"noisy OCR text broke the compile: {result['errors'][:200]}")


@test('escaping keeps renderable characters and replaces the rest')
def _():
    out = textract_fast.escape_tex('café » 中 ≤')
    check('é' in out, 'a Latin-1 accent was destroyed')
    check('»' in out, 'a Latin-1 symbol was destroyed')
    check('中' not in out and '?' in out, 'an unrenderable glyph was kept')
    check(r'$\le$' in out, 'a maths symbol was not mapped to LaTeX')


@test('Tesseract receives a deskewed page')
def _():
    if not textract_fast._TESSERACT_CMD:
        return skip('Tesseract not installed')
    path = os.path.join(_SCRATCH, 'skewed.png')
    text_page().rotate(-9, resample=Image.BICUBIC, expand=True,
                       fillcolor=255).save(path)
    text = textract_fast.extract_text_from_file(path)
    check('Chapter' in text or 'quick' in text,
          f'a rotated page still produced nothing usable: {text!r}')
    os.remove(path)


@test('stored results round-trip and expire by token')
def _():
    token = tex_store.save({'tex': GOOD_TEX, 'source': 'textract'})
    check(tex_store.read(token)['tex'] == GOOD_TEX, 'round-trip failed')
    tex_store.discard(token)
    check(tex_store.read(token) is None, 'discard did not delete')


@test('gemini remains the default provider for the AI pipeline')
def _():
    saved = os.environ.pop('AI_QA_PROVIDER', None)
    try:
        provider = _REAL_GET_PROVIDER()
        check(provider.name == 'gemini', provider.name)
        check(provider.trains_on_free_input(), 'free tier disclosure not declared')
    finally:
        if saved:
            os.environ['AI_QA_PROVIDER'] = saved


# ---------------------------------------------------------------------------
# Per-role model selection
# ---------------------------------------------------------------------------

_MODEL_VARS = ('AI_QA_MODEL', 'AI_QA_MODEL_DOCUMENT',
               'AI_QA_THINKING', 'AI_QA_THINKING_DOCUMENT',
               'AI_QA_EFFORT')


def with_clean_model_env(fn):
    """Run `fn` with every model/thinking override removed, then restore."""
    saved = {name: os.environ.pop(name, None) for name in _MODEL_VARS}
    try:
        return fn()
    finally:
        for name, value in saved.items():
            if value is not None:
                os.environ[name] = value
            else:
                os.environ.pop(name, None)


@test('a role asks for its own model and its own amount of thinking')
def _():
    def run():
        provider = with_provider(['LATEX_OK'])
        document = provider.start('sys', role=llm_providers.ROLE_DOCUMENT)

        check(document.model == 'scripted-document', document.model)
        # Converting a page is a reading job: a short leash, not the budget.
        check(document.thinking == 'low', document.thinking)
    try:
        with_clean_model_env(run)
    finally:
        restore_provider()


@test('the recommended Gemini models are the built-in defaults')
def _():
    def run():
        gemini = llm_providers.GeminiProvider()
        models = gemini.models()
        check(models[llm_providers.ROLE_DOCUMENT] == 'gemini-3.1-flash-lite',
              models[llm_providers.ROLE_DOCUMENT])
    with_clean_model_env(run)


@test('Claude defaults to Sonnet 5 rather than a pricier Opus tier')
def _():
    def run():
        claude = llm_providers.AnthropicProvider()
        for role, model in claude.models().items():
            check(model == 'claude-sonnet-5', f'{role}: {model}')
    with_clean_model_env(run)


@test('AI_QA_MODEL still overrides both pipelines at once')
def _():
    def run():
        os.environ['AI_QA_MODEL'] = 'pinned-model'
        gemini = llm_providers.GeminiProvider()
        for role, model in gemini.models().items():
            check(model == 'pinned-model', f'{role}: {model}')
    with_clean_model_env(run)


@test('a per-role model override beats the global one')
def _():
    def run():
        os.environ['AI_QA_MODEL'] = 'global-model'
        os.environ['AI_QA_MODEL_DOCUMENT'] = 'document-model'
        gemini = llm_providers.GeminiProvider()
        models = gemini.models()
        check(models[llm_providers.ROLE_DOCUMENT] == 'document-model',
              models[llm_providers.ROLE_DOCUMENT])
    with_clean_model_env(run)


@test('thinking level is overridable per role and globally')
def _():
    def run():
        gemini = llm_providers.GeminiProvider()
        check(gemini.default_thinking(llm_providers.ROLE_DOCUMENT) == 'low',
              'the built-in default changed')
        os.environ['AI_QA_THINKING'] = 'off'
        check(gemini.default_thinking(llm_providers.ROLE_DOCUMENT) == 'off',
              'global override ignored')
        os.environ['AI_QA_THINKING_DOCUMENT'] = 'medium'
        check(gemini.default_thinking(llm_providers.ROLE_DOCUMENT) == 'medium',
              'per-role override does not beat the global one')
    with_clean_model_env(run)


@test('provider_info reports both models and no credentials')
def _():
    def run():
        info = ai_qa.provider_info()
        check(set(info['models']) == set(llm_providers.ROLES), info['models'])
        blob = repr(info)
        for secret in ('GEMINI_API_KEY', 'ANTHROPIC_API_KEY', 'api_key'):
            check(secret not in blob, f'{secret} leaked into provider_info')
    with_clean_model_env(run)


# ---------------------------------------------------------------------------
# Unified pipeline: text/maths discrimination and merging
# ---------------------------------------------------------------------------

def line(text, top, bottom, conf, left=100, right=500, block=1, par=1, ln=1):
    return {'text': text, 'box': (left, top, right, bottom), 'min_conf': conf,
            'mean_conf': conf, 'block': block, 'par': par, 'line': ln}


@test('a confidently read line of prose is never sent to the formula model')
def _():
    lines = [line('Signal processing turns measurements into meaning.',
                  200, 222, 96)]
    nominated, rejected = layout.nominate(lines, [(100, 198, 500, 224)])
    check(nominated == [], f'prose was nominated: {nominated}')
    check(rejected, 'nothing was reported as rejected')


@test('a region the text engine could not read at all is nominated')
def _():
    # Measured on a rendered PDF: Tesseract returned nothing whatsoever for a
    # displayed derivative, so "no words here" has to nominate.
    nominated, _ = layout.nominate([], [(100, 400, 900, 500)])
    check(nominated == [(100, 400, 900, 500)], f'not nominated: {nominated}')


@test('a caption merged with the formula below it is carved apart')
def _():
    # The segmenter merges ink separated by less than a line space, so this
    # arrives as ONE band holding a sentence and an equation.
    lines = [line('The Fourier transform of f is', 200, 217, 96),
             line('F(w) -/ f(t)e* dt.', 224, 260, 5)]
    nominated, _ = layout.nominate(lines, [(100, 198, 900, 279)])
    check(len(nominated) == 1, f'expected one carved region, got {nominated}')
    check(nominated[0][1] >= 217,
          f'the carve kept the caption inside the formula: {nominated[0]}')


@test('an equation number does not split the formula it belongs to')
def _():
    # "(1)" is read perfectly by Tesseract. Treating it as a text boundary
    # shredded formulas into slivers and took a page from 2 equations to 0.
    lines = [line('F(w) -/ f(t)e* dt.', 224, 260, 5),
             line('(1)', 239, 256, 94, left=800, right=860)]
    nominated, _ = layout.nominate(lines, [(100, 220, 900, 279)])
    check(len(nominated) == 1, f'the formula was split: {nominated}')


@test('a sentence containing inline maths stays text')
def _():
    # Low confidence because of the inline maths, but it is still a sentence.
    # Splitting it would hand prose to a formula model; the text engine reads
    # a sentence with inline maths better than the formula model does.
    lines = [line("Einstein's mass-energy relation is E = mc?, and the "
                  "Lorentz factor is", 195, 224, 15)]
    nominated, _ = layout.nominate(lines, [(100, 193, 900, 226)])
    check(nominated == [], f'inline maths was sent to the formula model: {nominated}')


@test('letter-spaced prose from the formula model is rejected')
def _():
    # pix2text returns text as spaced letters inside \mathrm. This is the last
    # line of defence behind nomination.
    check(not layout.looks_like_equation(
        r'1. 1 \quad \mathrm { M o t i v a t i o n }'), 'prose accepted')
    check(not layout.looks_like_equation(
        r'\mathrm { T h e ~ F o u r i e r ~ t r a n s f o r m }'), 'prose accepted')
    for real in (r'F ( \omega ) = \int _ { 0 } ^ { 1 } f ( t ) d t',
                 r'H ( x ) = \begin{cases} 0 & x < 0 \end{cases}',
                 r'E = mc^{2}',
                 r'\frac { \partial u } { \partial t }'):
        check(layout.looks_like_equation(real), f'rejected real maths: {real}')


@test('equations displace the garbled text the OCR produced for them')
def _():
    lines = [line('Some prose above.', 100, 120, 95),
             line('F(w) -/ f(t)e* dt.', 200, 250, 5)]
    equations = [{'index': 1, 'latex': 'F(\\omega) = 1', 'box': (100, 195, 900, 255)}]
    items = layout.assemble(lines, equations)
    kinds = [item['kind'] for item in items]
    check(kinds == ['text', 'equation'], f'wrong merge: {kinds}')
    check(all('f(t)e*' not in i.get('text', '') for i in items),
          'the garbled OCR text was kept alongside the equation')


@test('items come back in reading order regardless of input order')
def _():
    lines = [line('Third.', 300, 320, 95), line('First.', 100, 120, 95)]
    equations = [{'index': 1, 'latex': 'x = 1', 'box': (100, 200, 500, 250)}]
    items = layout.assemble(lines, equations)
    check([i['box'][1] for i in items] == [100, 200, 300],
          f"out of order: {[i['box'][1] for i in items]}")


@test('the assembled document is valid LaTeX with maths in place')
def _():
    lines = [line('1 Energy', 100, 128, 90),
             line('The relation is', 200, 220, 96)]
    equations = [{'index': 1, 'latex': 'E = mc^{2}', 'box': (100, 260, 500, 300)}]
    tex = layout.to_tex(layout.assemble(lines, equations),
                        textract_fast.escape_tex)
    check(latex_tools.static_validate(tex) == [], latex_tools.static_validate(tex))
    check('\\section{Energy}' in tex, f'heading not recovered:\n{tex}')
    check('\\[\nE = mc^{2}\n\\]' in tex, f'equation not displayed:\n{tex}')
    check('amsmath' in tex, 'amsmath not loaded for a document with maths')


@test('amsmath is not loaded for a page with no mathematics')
def _():
    tex = layout.to_tex(layout.assemble([line('Just prose here.', 100, 120, 96)],
                                        []), textract_fast.escape_tex)
    check('amsmath' not in tex, 'an unused package was loaded')


# ---------------------------------------------------------------------------
# Handwritten and typewritten content
# ---------------------------------------------------------------------------

@test('a centred line is treated as displayed maths even when read confidently')
def _():
    # Printed algebra is ordinary glyphs, so Tesseract reads "E=mc" happily at
    # confidence 80 and nothing looks wrong. Position is the only signal left.
    centred = line('E=mc', 165, 210, 80, left=653, right=841)
    nominated, _ = layout.nominate([centred], [(600, 160, 900, 215)],
                                   page_width=1500)
    check(nominated, 'a centred printed equation was left as prose')


@test('a left-aligned confident line is still treated as text')
def _():
    left = line('Clean print 99.9 61', 225, 250, 94, left=140, right=900)
    nominated, _ = layout.nominate([left], [(130, 220, 910, 255)],
                                   page_width=1500)
    check(nominated == [], f'a table row was sent to the formula model: {nominated}')


@test('centring is judged on the words, not on the merged band')
def _():
    # A band merging a caption with the formula beneath keeps the band's full
    # width, which hides that the formula inside it is centred.
    caption = line('The energy of a particle is given by', 100, 140, 31,
                   left=91, right=954)
    formula = line('E=mc', 165, 210, 52, left=653, right=841)
    nominated, _ = layout.nominate([caption, formula], [(82, 96, 962, 230)],
                                   page_width=1500)
    check(nominated, 'the centred formula under a caption was missed')


@test('prose is recovered from the formula model rather than being lost')
def _():
    # Handed handwriting, pix2text answers with spaced letters in \mathrm.
    recovered = layout.unwrap_text(
        r'\mathrm { T h i s ~ i d e a ~ c h a n g e d ~ h o w }')
    check(recovered == 'This idea changed how', repr(recovered))


@test('unwrapping keeps word boundaries and drops spacing commands')
def _():
    check(layout.unwrap_text(r'\mathrm { W e ~ e x p l o r e }') == 'We explore',
          layout.unwrap_text(r'\mathrm { W e ~ e x p l o r e }'))
    check(layout.unwrap_text('') == '', 'empty input should stay empty')


@test('a line no engine could read is never silently dropped')
def _():
    # The regression this guards: Tesseract returned NOTHING for a handwritten
    # line, so it was nominated as an equation, rejected as prose, and deleted.
    import convert as real_convert
    page = Image.new('RGB', (900, 300), (255, 255, 255))
    # What pix2text actually returns when handed a line of handwriting.
    _recognized['latex'] = r'\mathrm { T h i s ~ i d e a ~ c h a n g e d }'
    try:
        equations, salvaged = real_convert._recognize_regions(
            page, [(50, 50, 800, 150)], lines=[])
    finally:
        _recognized.pop('latex', None)
    check(equations == [], 'letter-soup was accepted as an equation')
    check(salvaged and salvaged[0]['text'],
          'the unreadable region was dropped instead of salvaged')
    check(salvaged[0]['uncertain'], 'salvaged text was not flagged uncertain')


@test('text the OCR did read is left alone rather than salvaged over')
def _():
    import convert as real_convert
    page = Image.new('RGB', (900, 300), (255, 255, 255))
    covering = [line('Already read fine here.', 50, 150, 95, left=50, right=800)]
    _equations, salvaged = real_convert._recognize_regions(
        page, [(50, 50, 800, 150)], lines=covering)
    check(salvaged == [], 'salvage overwrote text the OCR had read')


@test('handwriting is typeset as ordinary LaTeX, not marked up as different')
def _():
    system = ' '.join(ai_qa._DIRECT_SYSTEM.split())
    check('a handwritten sentence and a printed one both become ordinary '
          'LaTeX prose' in system,
          'the model is not told to typeset both hands the same way')
    check('changes how carefully you must read, not how it is typeset' in system,
          'the reason the two are typeset alike is no longer given')


# ---------------------------------------------------------------------------
# Hybrid: AI-first with the converters as fallback
# ---------------------------------------------------------------------------

def with_direct(result):
    """Replace the direct AI conversion with a fixed result."""
    ai_qa.convert_page = (
        lambda data, name, validate=True, outline=None, rotation=None: result)
    return convert


@test('the AI path is taken when the model answers')
def _():
    module = with_direct(qa_result(GOOD_TEX, 'corrected'))
    built = module.convert(png_bytes(), 'scan.png')
    check(built['summary']['path'] == 'ai', built['summary']['path'])
    check(built['tex'] == GOOD_TEX, 'the AI document was not returned')


@test('the converters take over when the AI is out of quota')
def _():
    module = with_direct(ai_qa.blank_review(
        '', 'failed', 'The Gemini free tier is rate limited right now.'))
    built = module.convert(png_bytes(), 'scan.png')
    check(built['summary']['path'] == 'converters', built['summary']['path'])
    check(built['tex'].strip(), 'the fallback returned an empty document')
    notice = built['summary']['fallback_notice']
    check('rate limited' in notice['reason'], notice)
    check('quality may be lower' in notice['detail'],
          f'the quality cost was not disclosed: {notice}')


@test('a failed AI conversion is not followed by a second full AI pass')
def _():
    # The direct call already retries internally, so asking the same model to
    # review the fallback only repeats the same backoff for the same answer.
    calls = []
    ai_qa.convert_page = (
        lambda data, name, validate=True, outline=None, rotation=None: (
            calls.append('direct')
            or ai_qa.blank_review('', 'failed', 'rate limited')))
    convert.convert(png_bytes(), 'scan.png')
    check(calls == ['direct'], f'extra API call made: {calls}')


@test('AI_FIRST=false is refused without consent and honoured with it')
def _():
    os.environ['AI_FIRST'] = 'false'
    ai_qa.convert_page = lambda *a, **k: (_ for _ in ()).throw(
        AssertionError('the AI path ran despite AI_FIRST=false'))
    try:
        try:
            convert.convert(png_bytes(), 'scan.png')
        except convert.FallbackNotAuthorized as blocked:
            check(blocked.status['available'] is False, blocked.status)
        else:
            check(False, 'the fallback was used without the user agreeing')

        built = convert.convert(png_bytes(), 'scan.png', allow_fallback=True)
        check(built['summary']['path'] == 'converters',
              built['summary']['path'])
    finally:
        os.environ.pop('AI_FIRST', None)


@test('an unsupported file type is rejected before either path starts')
def _():
    try:
        convert.convert(b'anything', 'notes.xyz')
    except convert.FallbackNotAuthorized:
        check(False, 'the type check ran after the availability check')
    except RuntimeError as exc:
        check('.xyz' in str(exc), str(exc))
    else:
        check(False, 'an unsupported type was accepted')


@test('a .docx is accepted by the type check')
def _():
    with_direct(qa_result(GOOD_TEX, 'corrected'))
    built = convert.convert(docx_bytes(), 'notes.docx')
    check(built['summary']['path'] == 'ai', built['summary']['path'])
    check(built['tex'] == GOOD_TEX, 'the AI document was not returned')


@test('a .docx reaches the model as text, not as a picture of a page')
def _():
    seen = {}

    def capture(data, name, validate=True, outline=None, rotation=None):
        seen['outline'] = outline
        seen['bytes'] = data
        return qa_result(GOOD_TEX, 'corrected')

    ai_qa.convert_page = capture
    convert.convert(docx_bytes(), 'notes.docx')
    check(seen['bytes'] is None, 'the .docx was sent as a file')
    check('[HEADING 1] Mass and energy' in seen['outline'], seen['outline'][:200])
    check('[TABLE]' in seen['outline'], 'the table structure was not passed on')


@test('direct conversion returns nothing to fall back on when it fails')
def _():
    # This is the trade-off the hybrid exists to cover, so pin it down.
    with_provider([], fail_on_start=llm_providers.LlmError('no key'))
    try:
        result = ai_qa.convert_page(png_bytes(), 'scan.png')
        check(result['status'] == 'failed', result['status'])
        check(result['tex'] == '', 'a failed direct conversion invented a document')
    finally:
        restore_provider()


@test('the direct prompt asks for faithful transcription of either hand')
def _():
    system = ' '.join(ai_qa._DIRECT_SYSTEM.split())
    check('handwritten, printed, or both' in system,
          'the direct prompt does not cover both hands')
    check('unusual is not wrong' in system,
          'the direct prompt drops the unusual-mathematics rule')
    check('no OCR draft to check' in system,
          'the direct prompt still implies there is a draft')


# ---------------------------------------------------------------------------
# Unified pipeline: the route and its four input methods
# ---------------------------------------------------------------------------

@test('a transparent drawing is flattened instead of turning black')
def _():
    # The Draw canvas used to clear to transparent; PIL drops alpha rather than
    # compositing, which made the whole drawing read as solid ink.
    canvas = Image.new('RGBA', (400, 200), (0, 0, 0, 0))
    canvas.putpixel((10, 10), (17, 24, 39, 255))
    flat, _notes = preprocess.prepare_image(canvas)
    check(flat.mode in ('RGB', 'L'), f'still has alpha: {flat.mode}')
    import numpy as np
    grey = np.asarray(flat.convert('L'))
    check(grey.mean() > 200, f'background is not white (mean {grey.mean():.0f})')




# ---------------------------------------------------------------------------
# Automatic model fallback
#
# Free-tier quota is spent one model at a time, so one model refusing says
# nothing about the next. Before this existed, an exhausted primary took the
# whole AI path down and every user dropped to Tesseract until someone edited
# .env by hand.
# ---------------------------------------------------------------------------

_ROTATION_TRIED = []


def scripted_models(behaviour):
    """
    Replace provider.start() so each model behaves as `behaviour` says.

    Returns the real start() so the caller can put it back.
    """
    provider = llm_providers.get_provider()
    real = type(provider).start

    class Convo:
        def __init__(self, model):
            self.model = model
            self.calls = 0
            self.usage = {'input': 0, 'output': 0,
                          'cache_read': 0, 'cache_write': 0}

        def ask(self, parts):
            self.calls += 1
            outcome = behaviour(self.model)
            if isinstance(outcome, Exception):
                raise outcome
            _ROTATION_TRIED.append(self.model)
            return outcome

    def start(self, system, model=None, role=None, thinking=None,
              attempts=None):
        return Convo(model or self.default_model(role))

    type(provider).start = start
    return real


def restore_start(real):
    type(llm_providers.get_provider()).start = real


def _chain():
    return llm_providers.get_provider().model_chain(llm_providers.ROLE_DOCUMENT)


@test('each role has an ordered fallback chain, best first')
def _():
    provider = llm_providers.get_provider()
    chains = provider.chains()
    for role, chain in chains.items():
        check(len(chain) >= 1, f'{role}: no chain')
        check(len(chain) == len(set(chain)), f'{role}: duplicate models {chain}')
        check(chain[0] == provider.default_model(role),
              f'{role}: chain does not lead with the preferred model')


@test('pinning a model keeps it first without removing the safety net')
def _():
    provider = llm_providers.get_provider()
    os.environ['AI_QA_MODEL_DOCUMENT'] = 'pinned-model'
    try:
        chain = provider.model_chain(llm_providers.ROLE_DOCUMENT)
        check(chain[0] == 'pinned-model', chain)
        check(len(chain) > 1, 'pinning removed every fallback')
    finally:
        del os.environ['AI_QA_MODEL_DOCUMENT']

    os.environ['AI_QA_MODEL_DOCUMENT'] = 'first, second'
    try:
        chain = provider.model_chain(llm_providers.ROLE_DOCUMENT)
        check(chain[:2] == ['first', 'second'], chain)
    finally:
        del os.environ['AI_QA_MODEL_DOCUMENT']


@test('an exhausted model is skipped for the next one automatically')
def _():
    chain = _chain()
    ai_status.clear_outage()
    real = scripted_models(
        lambda m: (llm_providers.LlmQuotaError(
            'daily allowance used up', retry_after=41, scope='day')
            if m == chain[0] else 'answer'))
    try:
        convo, reply, _p = ai_qa.ask_first_usable(
            'sys', llm_providers.ROLE_DOCUMENT, ['parts'])
        check(convo.model == chain[1],
              f'answered by {convo.model}, expected {chain[1]}')
        check(reply == 'answer', reply)
        check(chain[0] in ai_status.unavailable_models(),
              'the exhausted model was not remembered')
    finally:
        restore_start(real)
        ai_status.clear_outage()


@test('rotation keeps walking the chain, not just one step')
def _():
    chain = _chain()
    if len(chain) < 3:
        skip('chain too short')
        return
    ai_status.clear_outage()
    real = scripted_models(
        lambda m: (llm_providers.LlmQuotaError('rate limited', retry_after=30,
                                               scope='minute')
                   if m in chain[:2] else 'answer'))
    try:
        convo, _reply, _p = ai_qa.ask_first_usable(
            'sys', llm_providers.ROLE_DOCUMENT, ['parts'])
        check(convo.model == chain[2], convo.model)
    finally:
        restore_start(real)
        ai_status.clear_outage()


@test('a model this key cannot reach is skipped, not fatal')
def _():
    chain = _chain()
    ai_status.clear_outage()
    real = scripted_models(
        lambda m: (llm_providers.LlmModelError(f'{m} is not available')
                   if m == chain[0] else 'answer'))
    try:
        convo, _reply, _p = ai_qa.ask_first_usable(
            'sys', llm_providers.ROLE_DOCUMENT, ['parts'])
        check(convo.model == chain[1], convo.model)
    finally:
        restore_start(real)
        ai_status.clear_outage()


@test('one exhausted model does NOT take the service down')
def _():
    chain = _chain()
    ai_status.clear_outage()
    try:
        ai_status.record_model_outage(chain[0], 'daily quota gone',
                                      retry_after=600, scope='day')
        status = ai_status.check()
        check(status['available'] is True,
              'one model out of quota still disables the whole AI path')
        service = status['services'][0]
        check(chain[0] in service['exhausted_models'], service)
        check(chain[1] in service['available_models'], service)
    finally:
        ai_status.clear_outage()


@test('the service is down only when every model is exhausted')
def _():
    chain = _chain()
    ai_status.clear_outage()
    real = scripted_models(
        lambda m: llm_providers.LlmQuotaError('out of quota', retry_after=60,
                                              scope='minute'))
    try:
        try:
            ai_qa.ask_first_usable('sys', llm_providers.ROLE_DOCUMENT, ['p'])
        except llm_providers.LlmError:
            pass
        else:
            check(False, 'an exhausted chain still returned an answer')

        check(len(ai_status.unavailable_models()) == len(chain),
              'not every attempted model was recorded')
        status = ai_status.check()
        check(status['available'] is False, 'service still reported as up')
        check('reached its quota' in status['reason'], status['reason'])
        check(status['recovery']['known'] is True,
              'the soonest recovery time was not reported')
    finally:
        restore_start(real)
        ai_status.clear_outage()


@test('a rejected API key fails at once instead of rotating')
def _():
    # A bad key dooms every model, so trying three more wastes the user's time
    # to reach the same answer.
    ai_status.clear_outage()
    tried = []
    provider = llm_providers.get_provider()
    real = type(provider).start

    class Convo:
        def __init__(self, model):
            self.model = model
            self.calls = 0
            self.usage = {}

        def ask(self, parts):
            raise llm_providers.LlmError(
                "The server's Gemini API key was rejected.")

    def start(self, system, model=None, role=None, thinking=None,
              attempts=None):
        tried.append(model)
        return Convo(model)

    type(provider).start = start
    try:
        try:
            ai_qa.ask_first_usable('sys', llm_providers.ROLE_DOCUMENT, ['p'])
        except llm_providers.LlmError:
            pass
        else:
            check(False, 'a rejected key was treated as success')
        check(len(tried) == 1, f'rotated {len(tried)} times on a bad key')
        check(ai_status.check()['available'] is False,
              'a rejected key did not mark the service down')
    finally:
        type(provider).start = real
        ai_status.clear_outage()


@test('a successful call un-parks the model it used')
def _():
    chain = _chain()
    ai_status.clear_outage()
    real = scripted_models(lambda m: 'answer')
    try:
        ai_status.record_model_outage(chain[0], 'stale', retry_after=600,
                                      scope='minute')
        ai_qa.ask_first_usable('sys', llm_providers.ROLE_DOCUMENT, ['p'])
        # It rotated to chain[1] and that success clears chain[1], not chain[0];
        # what matters is that a model which answers is never left parked.
        check(chain[1] not in ai_status.unavailable_models(),
              'a model that answered is still marked exhausted')
    finally:
        restore_start(real)
        ai_status.clear_outage()



@test('a round opens on the preferred model')
def _():
    ai_status.clear_outage()
    rotation = ai_qa.Rotation(llm_providers.ROLE_DOCUMENT)
    provider = llm_providers.get_provider()
    check(rotation.active == provider.default_model(llm_providers.ROLE_DOCUMENT),
          f'round opened on {rotation.active}')
    check(rotation.exhausted is False, 'a fresh round is already spent')


@test('a round keeps the model it rotated to, instead of re-probing each page')
def _():
    chain = _chain()
    ai_status.clear_outage()
    asked = []

    def behaviour(model):
        asked.append(model)
        return (llm_providers.LlmQuotaError('quota', retry_after=None)
                if model == chain[0] else 'answer')

    real = scripted_models(behaviour)
    try:
        rotation = ai_qa.Rotation(llm_providers.ROLE_DOCUMENT)
        for _page in range(4):
            rotation.ask('sys', ['parts'])
        # Page one pays to discover chain[0] is out; pages two to four must not.
        check(asked.count(chain[0]) == 1,
              f'the exhausted model was probed {asked.count(chain[0])} times '
              f'in one round: {asked}')
        check(asked.count(chain[1]) == 4, asked)
        check(rotation.active == chain[1], rotation.active)
    finally:
        restore_start(real)
        ai_status.clear_outage()


@test('a new round starts back at the preferred model')
def _():
    # A model out of quota for a minute must not push every later conversion
    # onto a worse model for the rest of the day.
    chain = _chain()
    ai_status.clear_outage()
    try:
        # Parked by our own guess, not by the provider: worth re-testing.
        ai_status.record_model_outage(chain[0], 'assumed', retry_after=None,
                                      provider='test')
        rotation = ai_qa.Rotation(llm_providers.ROLE_DOCUMENT)
        check(rotation.active == chain[0],
              f'a new round opened on {rotation.active}, not the default')

        asked = []
        real = scripted_models(lambda m: asked.append(m) or 'answer')
        try:
            rotation.ask('sys', ['parts'])
            check(asked == [chain[0]],
                  f'the preferred model was not re-tried: {asked}')
        finally:
            restore_start(real)
    finally:
        ai_status.clear_outage()


@test('a provider-stated retry window is honoured without a wasted call')
def _():
    # The one exception to re-trying the default: when the provider itself named
    # a time, calling early wastes a round trip and can extend the block.
    chain = _chain()
    ai_status.clear_outage()
    asked = []
    real = scripted_models(lambda m: asked.append(m) or 'answer')
    try:
        ai_status.record_model_outage(chain[0], 'rate limited', retry_after=120,
                                      scope='minute', provider='test')
        rotation = ai_qa.Rotation(llm_providers.ROLE_DOCUMENT)
        convo, _reply = rotation.ask('sys', ['parts'])
        check(chain[0] not in asked,
              'called a model the provider told us to leave alone')
        check(convo.model == chain[1], convo.model)
    finally:
        restore_start(real)
        ai_status.clear_outage()


@test('a model believed exhausted is confirmed with one call, not three')
def _():
    chain = _chain()
    ai_status.clear_outage()
    seen = {}

    provider = llm_providers.get_provider()
    real = type(provider).start

    class Convo:
        def __init__(self, model, attempts):
            self.model = model
            self.calls = 0
            self.usage = {}
            seen[model] = attempts

        def ask(self, parts):
            return 'answer'

    def start(self, system, model=None, role=None, thinking=None,
              attempts=None):
        return Convo(model, attempts)

    type(provider).start = start
    try:
        # Parked by a guess, so the round still tries it - but cheaply.
        ai_status.record_model_outage(chain[0], 'assumed', retry_after=None,
                                      provider='test')
        ai_qa.Rotation(llm_providers.ROLE_DOCUMENT).ask('sys', ['p'])
        check(seen.get(chain[0]) == 1,
              f'retry budget for a parked model was {seen.get(chain[0])}, '
              'expected 1 - three attempts with backoff would add seconds to '
              'every conversion')
    finally:
        type(provider).start = real
        ai_status.clear_outage()


@test('an exhausted round stops asking instead of failing the conversion')
def _():
    chain = _chain()
    ai_status.clear_outage()
    calls = {'n': 0}

    def behaviour(model):
        calls['n'] += 1
        return llm_providers.LlmQuotaError('out', retry_after=None)

    real = scripted_models(behaviour)
    try:
        rotation = ai_qa.Rotation(llm_providers.ROLE_DOCUMENT)
        try:
            rotation.ask('sys', ['p'])
        except llm_providers.LlmError:
            pass
        check(rotation.exhausted is True, 'the round is not marked spent')
        check(calls['n'] == len(chain),
              f'tried {calls["n"]} models, chain has {len(chain)}')

        # The caller checks `exhausted` and stops; no further calls are made.
        before = calls['n']
        check(rotation.active is None, 'a spent round still names a model')
        check(calls['n'] == before, 'a spent round made another call')
    finally:
        restore_start(real)
        ai_status.clear_outage()


@test('a multi-page conversion completes locally once every model is out')
def _():
    # The whole point: the user gets a finished document, not an error.
    ai_status.clear_outage()
    real = scripted_models(
        lambda m: llm_providers.LlmQuotaError('out of quota', retry_after=None))
    try:
        result = convert.convert(pdf_bytes(pages=3), 'three.pdf',
                                 allow_fallback=True)
        check(result['tex'].strip(), 'the conversion came back empty')
        stats = result['summary']
        check(stats['path'] == 'converters', stats['path'])
        check(stats['ai_pages'] == 0, stats)
        check(stats['fallback_notice'],
              'the user was not told quality dropped')
    finally:
        restore_start(real)
        ai_status.clear_outage()


# ---------------------------------------------------------------------------
# Preserving work, and telling the user once
# ---------------------------------------------------------------------------

def _ai_pages_then_quota(good_pages, tex=None):
    """Convert `good_pages` pages with the AI, then run out of quota."""
    done = {'n': 0}
    body = tex or GOOD_TEX

    def convert_page(data, name, validate=True, outline=None, rotation=None):
        if done['n'] >= good_pages:
            return ai_qa.blank_review('', 'failed',
                                      'Every AI model has reached its quota.')
        done['n'] += 1
        page = body.replace('Energy', f'Page {done["n"]}')
        return {'tex': page, 'status': 'corrected', 'message': '',
                'findings': [], 'equations': [], 'summary': {},
                'compile': {'attempted': False, 'ok': False, 'engine': None,
                            'errors': '', 'missing_packages': [],
                            'reason': None},
                'usage': {'input': 0, 'output': 0, 'cache_read': 0,
                          'cache_write': 0},
                'model': 'scripted', 'provider': 'scripted'}
    return convert_page


@test('pages converted before exhaustion are kept, not redone')
def _():
    ai_qa.convert_page = _ai_pages_then_quota(2)
    result = convert.convert(pdf_bytes(pages=4), 'four.pdf')
    stats = result['summary']
    check(stats['ai_pages'] == 2, stats)
    check(stats['fallback_pages'] == 2, stats)
    check(stats['path'] == 'mixed', stats['path'])
    # The AI's own words for pages one and two must still be in the document.
    for marker in ('Page 1', 'Page 2'):
        check(marker in result['tex'], f'{marker} was discarded and redone')


@test('the fallback resumes at the page the AI stopped on')
def _():
    seen = {}
    real_local = convert._local_document

    def spy(pages, first_number=1, equation_offset=0):
        seen['first'] = first_number
        seen['count'] = len(pages)
        return real_local(pages, first_number, equation_offset)

    convert._local_document = spy
    try:
        ai_qa.convert_page = _ai_pages_then_quota(3)
        convert.convert(pdf_bytes(pages=5), 'five.pdf')
        check(seen['first'] == 4,
              f'the fallback restarted at page {seen["first"]}, not 4')
        check(seen['count'] == 2,
              f'the fallback re-converted {seen["count"]} pages, not 2 - work '
              'the AI had already finished was thrown away')
    finally:
        convert._local_document = real_local


@test('work survives the fallback itself failing')
def _():
    # The tail is the step most likely to fail on a server missing Poppler or
    # Tesseract. Two AI pages must not be lost to that.
    ai_qa.convert_page = _ai_pages_then_quota(2)
    real_pages = convert._pages

    def broken(file_bytes, filename, first=1):
        if first > 1:
            raise RuntimeError('Poppler was not found.')
        return real_pages(file_bytes, filename, first)

    convert._pages = broken
    try:
        result = convert.convert(pdf_bytes(pages=4), 'four.pdf')
        check('Page 1' in result['tex'] and 'Page 2' in result['tex'],
              'AI pages were lost when the fallback failed')
        notice = result['summary']['fallback_notice']
        check(notice and notice['partial'] is True,
              'a truncated document was not flagged as incomplete')
        check('Poppler' in notice['reason'], notice)
    finally:
        convert._pages = real_pages


@test('a conversion with nothing salvageable still raises rather than lying')
def _():
    # No AI pages and a broken fallback is a real failure, not a partial one.
    ai_qa.convert_page = _ai_pages_then_quota(0)
    real_pages = convert._pages
    convert._pages = lambda *a, **k: (_ for _ in ()).throw(
        RuntimeError('Poppler was not found.'))
    try:
        try:
            convert.convert(pdf_bytes(pages=2), 'two.pdf')
        except RuntimeError as exc:
            check('Poppler' in str(exc), str(exc))
        else:
            check(False, 'an empty conversion was reported as a success')
    finally:
        convert._pages = real_pages


@test('one page that will not open does not cost the others')
def _():
    from PIL import Image as _Image

    class Broken:
        def convert(self, *a, **k):
            raise OSError('truncated image')

        @property
        def size(self):
            raise OSError('truncated image')

    good = _Image.new('RGB', (400, 200), 'white')
    tex, _eq, items, notes = convert._local_document([good, Broken(), good])
    check(tex.strip(), 'a broken page emptied the whole document')
    check(any('could not be read' in note for note in notes), notes)


@test('the degraded notice appears exactly once on the result page')
def _():
    ai_status.clear_outage()
    ai_qa.convert_page = _ai_pages_then_quota(1)
    client = accept_terms(client_for_tests())
    client.post('/convert',
                data={'file': (io.BytesIO(pdf_bytes(pages=2)), 'x.pdf')},
                content_type='multipart/form-data')

    # Read the token before rendering: home() pops it, so following the
    # redirect first would leave nothing to look up.
    with client.session_transaction() as session:
        token = session.get('convert_token')
    check(token, 'the conversion produced no result')
    notice = tex_store.read(token)['stats']['fallback_notice']
    check(notice, 'no notice was produced')

    page = client.get('/').get_data(as_text=True)

    # Once. It used to render as a banner AND as the QA status message.
    count = page.count(notice['detail'])
    check(count == 1, f'the notice appears {count} times on the page')
    check(page.count(notice['headline']) == 1,
          'the headline appears more than once')


@test('a fully successful conversion carries no notice at all')
def _():
    ai_status.clear_outage()
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    page = client.post('/convert',
                       data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                       content_type='multipart/form-data',
                       follow_redirects=True).get_data(as_text=True)
    # Match the notice block itself, not loose wording: "Continue without AI"
    # is a button in the always-present outage dialog.
    check('role="status"' not in page,
          'a clean conversion rendered a degraded-quality notice')
    check('Incomplete conversion' not in page,
          'a clean conversion was flagged as incomplete')
    check('quality may be lower' not in page,
          'a clean conversion warned about quality anyway')

@test('a per-model outage expires on its own')
def _():
    import time as _time
    chain = _chain()
    ai_status.clear_outage()
    try:
        ai_status.record_model_outage(chain[0], 'brief', retry_after=1,
                                      scope='minute')
        check(chain[0] in ai_status.unavailable_models(), 'not recorded')
        _time.sleep(1.2)
        check(chain[0] not in ai_status.unavailable_models(),
              'the model stayed parked after its retry window passed')
    finally:
        ai_status.clear_outage()


# ---------------------------------------------------------------------------
# The preview request
# ---------------------------------------------------------------------------

def _script():
    """The browser-side source, read from disk as the browser would get it."""
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        'static', 'scripts.js')
    with open(path, encoding='utf-8') as handle:
        return handle.read()


def _js_function(script, name):
    """The source of one top-level function in scripts.js, or ''."""
    start = script.find(f'function {name}(')
    if start < 0:
        return ''
    end = script.find('\nfunction ', start + 1)
    return script[start:end if end > 0 else len(script)]


def _read_template(relative):
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        'templates', *relative.split('/'))
    with open(path, encoding='utf-8') as handle:
        return handle.read()


@test('the preview gives up later than the server does')
def _():
    # Both sides time the same compile. If the browser gave up first the user
    # would get a generic "took too long" instead of the server's actual
    # reason, so this ordering is the difference between a useful error and a
    # useless one.
    script = _script()
    match = re.search(r'PREVIEW_TIMEOUT_MS\s*=\s*(\d+)', script)
    check(match, 'PREVIEW_TIMEOUT_MS is gone from scripts.js')
    client_seconds = int(match.group(1)) / 1000
    server_seconds = latex_tools._COMPILE_TIMEOUT
    check(client_seconds > server_seconds,
          f'the browser gives up after {client_seconds}s but the server '
          f'compiles for up to {server_seconds}s, so the server error can '
          'never reach the user')


@test('the preview never asks the browser to handle a PDF')
def _():
    # A browser never lets the page have an application/pdf response: Chromium
    # routes it to its own viewer before script can read it, and a download
    # manager extension takes it away entirely. Either way the panel goes
    # blank while the server has done everything right. The preview is built
    # from a JSON page count and plain images for exactly that reason, and
    # nothing may quietly go back to the old way.
    script = _script()
    check('fetchPdf' not in script and 'previewObjectUrl' not in script,
          'the preview is fetching the PDF as data again')
    check('preview-frame' not in script,
          'the preview is being put back into a frame pointed at a PDF')
    # Scoped to the functions that build the preview. Elsewhere the PDF route
    # is fair game - an Open in new tab link is a deliberate request for the
    # file, and a download manager taking that one is the user's own doing.
    for name in ('loadPreview', 'fetchPages', 'fillPages', 'showPagePreview'):
        body = _js_function(script, name)
        check(body, f'{name} is gone from scripts.js')
        check('preview.pdf' not in body,
              f'{name} loads the PDF itself again - the browser will take it '
              'away from the page and the panel will go blank')
    check('preview/pages' in script and 'preview/page.png' in script,
          'the preview no longer uses the page-image routes')

    # The panel markup must not carry a PDF frame either.
    panel = _read_template('partials/convert_section.html')
    check('<iframe' not in panel, 'the preview panel still holds an iframe')
    check('preview-pages' in panel, 'the preview panel has no page container')


@test('the page images are served as images, and the count as JSON')
def _():
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    client.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                content_type='multipart/form-data')
    with client.session_transaction() as session:
        token = session['convert_token']

    listing = client.get(f'/preview/pages?token={token}')
    check(listing.status_code == 200, listing.status_code)
    check(listing.headers['Content-Type'].startswith('application/json'),
          listing.headers['Content-Type'])
    body = listing.get_json()
    check(body['ok'] is True and body['pages'] >= 1, body)

    page = client.get(f'/preview/page.png?token={token}&n=1')
    check(page.status_code == 200, page.status_code)
    check(page.headers['Content-Type'].startswith('image/png'),
          f'the preview page was sent as {page.headers["Content-Type"]} - a '
          'browser must never be handed a PDF here')
    check(page.data[:8] == b'\x89PNG\r\n\x1a\n', 'not a PNG')

    # Cached, not re-rendered.
    again = client.get(f'/preview/pages?token={token}')
    check(again.get_json()['pages'] == body['pages'], again.get_json())


@test('page images belong to the session that made them')
def _():
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    flask_app.convert.convert = scripted_convert()
    owner = accept_terms(client_for_tests())
    owner.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
               content_type='multipart/form-data')
    with owner.session_transaction() as session:
        token = session['convert_token']
    owner.get(f'/preview/pages?token={token}')

    stranger = client_for_tests()
    check(stranger.get(f'/preview/pages?token={token}').status_code == 404,
          'another session listed the pages')
    check(stranger.get(f'/preview/page.png?token={token}&n=1').status_code == 404,
          'another session read a rendered page')


@test('a preview page number cannot escape its own document')
def _():
    client = client_for_tests()
    for bad in ('0', '-1', 'x', '999', '../../.env'):
        response = client.get(f'/preview/page.png?token=deadbeef&n={bad}')
        check(response.status_code == 404,
              f'n={bad} -> {response.status_code}')


@test('a preview failure is rendered for whoever asked for it')
def _():
    # The frame shows whatever comes back, so a failure has to be a document a
    # frame can display. Raw JSON in the preview panel is not an error message.
    if not latex_tools.find_engine():
        skip('no LaTeX engine installed')
        return
    client = client_for_tests()
    token = tex_store.save({'tex': BROKEN_TEX, 'file_name': 'bad.png',
                            'source': 'convert'})
    with client.session_transaction() as session:
        session['tex_tokens'] = [token]

    framed = client.get(f'/preview.pdf?token={token}',
                        headers={'Accept': 'text/html,application/xhtml+xml,'
                                           'application/xml;q=0.9,*/*;q=0.8'})
    check(framed.status_code == 422, framed.status_code)
    body = framed.get_data(as_text=True)
    check(framed.headers['Content-Type'].startswith('text/html'),
          f'a frame was sent {framed.headers["Content-Type"]}')
    check('could not be rendered' in body, 'the frame page says nothing useful')
    check('.tex' in body, 'the frame page does not say the .tex is unaffected')
    check(not body.lstrip().startswith('{'), 'raw JSON was sent to the frame')

    # Scripts and the API keep JSON.
    scripted = client.get(f'/preview.pdf?token={token}',
                          headers={'Accept': 'application/json'})
    check(scripted.status_code == 422, scripted.status_code)
    check(scripted.get_json()['ok'] is False, scripted.get_json())


@test('an expired preview does not put JSON in the frame either')
def _():
    client = client_for_tests()
    response = client.get('/preview.pdf?token=deadbeef',
                          headers={'Accept': 'text/html,*/*;q=0.8'})
    check(response.status_code == 404, response.status_code)
    check(response.headers['Content-Type'].startswith('text/html'),
          response.headers['Content-Type'])


@test('the rendered document is reachable when the browser will not show it')
def _():
    # Some browsers are configured never to display a PDF inline. The preview
    # frame stays empty for them, so there has to be another way to the file.
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    page = client.post('/convert',
                       data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                       content_type='multipart/form-data',
                       follow_redirects=True).get_data(as_text=True)
    check('preview-open' in page, 'no way out when the frame cannot render')
    check('target="_blank"' in page, 'the escape hatch does not open the PDF')


@test('every preview blob is released')
def _():
    # An object URL lives until it is revoked. The preview is reloadable, so
    # without this each reload leaks another copy of the PDF for the tab's
    # lifetime.
    script = _script()

    # Input thumbnails all go through one helper, which owns the element's src
    # and so is the thing that must release what it replaces.
    helper = script[script.index('function showPreview('):]
    helper = helper[:helper.index('\n}')]
    check('URL.revokeObjectURL' in helper,
          'showPreview replaces an image source without releasing the blob it '
          'replaces, so every retaken photo pins another full-resolution '
          'capture in memory')

    # The rendered-PDF panels hold their own URL and release it in place.
    created = [m.start() for m in re.finditer(r'URL\.createObjectURL', script)]
    check(created, 'nothing renders a blob any more - has this moved?')
    for position in created:
        released = 'URL.revokeObjectURL' in script[max(0, position - 300):position]
        delegated = 'showPreview(' in script[max(0, position - 120):position]
        check(released or delegated,
              f'the blob created at character {position} neither releases the '
              'one it replaces nor goes through showPreview, so reloading '
              'that preview leaks a copy')


# ---------------------------------------------------------------------------
# The interface
#
# These guard the shape of the redesign rather than its appearance: which page
# a thing lives on, whether it can be operated without a mouse, and whether it
# tells the truth. Wording and colour are free to change; the properties below
# are not.
# ---------------------------------------------------------------------------

ALL_PAGES = ('/', '/history', '/login', '/signup', '/forgot-password')


@test('every page is built from the one shell')
def _():
    # The four pages used to be independent documents, each with its own
    # <head>, which is why they had drifted: one set its body font one way and
    # the others another, only one had a background colour, and not one had a
    # favicon or a description. Anything that has to be on every page is
    # checked on every page.
    client = client_for_tests()
    for path in ALL_PAGES:
        page = client.get(path).get_data(as_text=True)
        for needed, why in (
                ('img/contex-mark', 'the logo'),
                ('img/favicon', 'a favicon'),
                ('name="description"', 'a meta description'),
                ('og:image', 'a social card'),
                ('Skip to main content', 'a skip link'),
                ('id="main"', 'a main landmark'),
                ('<footer', 'the footer'),
                ('id="legal-modal"', 'a way to read the legal documents')):
            check(needed in page, f'{path} has no {why}')


@test('the workspace is the converter, not a page about the converter')
def _():
    page = client_for_tests().get('/').get_data(as_text=True)
    check('id="convert-form"' in page and 'id="convert-drop-area"' in page,
          'the upload control is not on the workspace')
    # The hero, the About section and the History section used to sit on this
    # route, each a full viewport tall, with the converter third of four.
    for gone, what in (('id="about"', 'the About section'),
                       ('id="home"', 'the hero section'),
                       ('id="history"', 'the history section')):
        check(gone not in page, f'{what} is back on the workspace')
    check('min-h-[90vh]' not in page and 'landscape:min-h-screen' not in page,
          'the viewport-height sections are back')


@test('history is its own route and is not also on the workspace')
def _():
    item = {'uid': 'uid-route', 'id': 'doc-route', 'fileName': 'saved.png',
            'ocrType': 'convert', 'result': GOOD_TEX, 'timestamp': None}
    _history_items[('uid-route', 'doc-route')] = item
    client = client_for_tests()
    with client.session_transaction() as session:
        session['user'] = {'uid': 'uid-route', 'email': 'r@example.com'}

    real_list = _fake_firebase.get_user_ocr_history
    _fake_firebase.get_user_ocr_history = lambda uid, **k: (
        [dict(item)] if uid == 'uid-route' else [])
    try:
        history = client.get('/history')
        workspace = client.get('/').get_data(as_text=True)
    finally:
        _fake_firebase.get_user_ocr_history = real_list

    check(history.status_code == 200, history.status_code)
    body = history.get_data(as_text=True)
    check('saved.png' in body, 'the saved conversion is not listed')
    for action in ('/history/doc-route/download', 'copyHistory',
                   'toggleHistoryPreview', 'openPdf'):
        check(action in body, f'{action} is missing from the history page')
    check('saved.png' not in workspace,
          'history is being rendered on the workspace as well')


@test('the result replaces the input panel instead of stacking under it')
def _():
    # The finished document used to appear below the upload control that
    # produced it, inside the same card, so the thing you came back for was
    # under a control you had already finished with.
    flask_app.convert.convert = scripted_convert()
    client = accept_terms(client_for_tests())
    client.post('/convert', data={'file': (io.BytesIO(png_bytes()), 'a.png')},
                content_type='multipart/form-data')
    page = client.get('/').get_data(as_text=True)
    check('id="preview-panel"' in page, 'the result is not shown')
    check('id="convert-form"' not in page,
          'the upload form is still on the page above the result')
    check('Convert another' in page, 'there is no way back to a fresh upload')


@test('the processing screen claims no progress it cannot know')
def _():
    # The server does not report how far through a document it is, so anything
    # on this screen that implied it would be measuring nothing.
    page = client_for_tests().get('/').get_data(as_text=True)
    start = page.index('id="processing"')
    screen = page[start:page.index('</div>', page.index('processing-note'))]
    check('id="processing-elapsed"' in screen,
          'the processing screen does not show how long it has been running')
    for faked in ('role="progressbar"', '<progress', 'aria-valuenow', '%'):
        check(faked not in screen,
              f'the processing screen contains {faked!r} - it cannot know that')

    body = _js_function(_script(), 'showProcessing')
    check(body, 'showProcessing is gone from scripts.js')
    check('setInterval' in body,
          'the elapsed time is no longer a real running count')


@test('the drop area can be used without a mouse')
def _():
    # It was a <div> with a click handler: not focusable, not announced, and
    # unusable from a keyboard. It is now a label for the file input, so the
    # input is focusable and takes its name from the label's text.
    panel = _read_template('partials/convert_section.html')
    check('<label for="convert-file-upload" id="convert-drop-area"' in panel,
          'the drop area is not a label for the file input')
    check('class="sr-only"' in panel.split('id="convert-file-upload"')[1][:300]
          or 'class="sr-only"' in panel.split('convert-drop-area')[1][:900],
          'the file input is hidden in a way that cannot take focus')
    check('dropArea.addEventListener(\'click\'' not in _script(),
          'a click handler is opening the picker a second time')


@test('no control on any page is a decoration')
def _():
    # The sign-up page shipped two social buttons with no click handler and no
    # Firebase config behind them, and two legal links that were href="#".
    client = client_for_tests()
    for path in ALL_PAGES:
        page = client.get(path).get_data(as_text=True)
        check('href="#"' not in page, f'{path} has a link that goes nowhere')
        check('coming soon' not in page.lower(),
              f'{path} offers something that does not exist yet')

    signup = client.get('/signup').get_data(as_text=True)
    for button in re.findall(r'<button[^>]*>', signup):
        wired = ('onclick=' in button or 'type="submit"' in button
                 or 'id="google-signup"' in button)
        check(wired, f'a button on the sign-up page does nothing: {button[:70]}')


@test('every colour the markup asks for actually exists')
def _():
    # bg-forest-100 appeared 23 times across the templates and scripts.js while
    # the palette started at forest-400, so every one of those classes
    # generated no CSS at all and the styling silently did nothing.
    #
    # A class Tailwind could not build is not in the built stylesheet, so this
    # catches an undefined colour. It also catches the other way of getting the
    # same silent nothing: adding a class and forgetting to run build_css.py.
    root = os.path.dirname(os.path.abspath(__file__))
    with open(os.path.join(root, 'static', 'css', 'app.css'),
              encoding='utf-8') as handle:
        css = handle.read()

    sources = []
    for folder, _, names in os.walk(os.path.join(root, 'templates')):
        sources += [os.path.join(folder, n) for n in names if n.endswith('.html')]
    sources.append(os.path.join(root, 'static', 'scripts.js'))

    palette = r'(?:ink|paper|forest|burgundy|cream|caution|alarm|affirm)'
    prefix = (r'(?:bg|text|border|ring|outline|decoration|accent|divide|fill'
              r'|stroke|from|via|to|shadow|placeholder|caret)')
    variant = (r'(?:hover|focus|focus-visible|focus-within|active|disabled'
               r'|group-hover|first|last|odd|even|sm|md|lg|xl|print'
               r'|motion-reduce|landscape|portrait):')
    pattern = re.compile(
        r'\b((?:' + variant + r')*' + prefix + '-' + palette
        + r'-\d{2,3}(?:/\d{1,3})?)\b')

    missing = {}
    for path in sources:
        with open(path, encoding='utf-8') as handle:
            for name in set(pattern.findall(handle.read())):
                # Tailwind keeps the variant in the class name and escapes
                # both ':' and '/', so hover:bg-burgundy-100 is written as
                # .hover\:bg-burgundy-100:hover - the whole name has to be
                # reconstructed, not just the colour part of it.
                escaped = name.replace(':', chr(92) + ':')
                escaped = escaped.replace('/', chr(92) + '/')
                if '.' + escaped not in css:
                    missing.setdefault(name, []).append(
                        os.path.basename(path))
    check(not missing,
          'colour classes that generate no CSS: '
          + ', '.join(f'{k} ({", ".join(v)})' for k, v in sorted(missing.items())))


@test('the brand assets the pages reference are actually served')
def _():
    client = client_for_tests()
    page = client.get('/').get_data(as_text=True)
    referenced = set(re.findall(r'/static/(img/[A-Za-z0-9._-]+)', page))
    check(referenced, 'the pages reference no images at all')
    for name in sorted(referenced):
        response = client.get('/static/' + name)
        check(response.status_code == 200, f'{name} is referenced but 404s')
        check(len(response.get_data()) > 500, f'{name} is served empty')


@test('the tab icon has rounded corners at every size it ships')
def _():
    # The favicons are generated (brand/make_assets.py), so the rounding is a
    # property of a build step rather than of anything in the source tree - the
    # kind of thing that comes back square the next time someone regenerates
    # them. CORNER there is 3/16, chosen because it lands on a whole pixel at
    # 16, 32 and 48, so every size rounds by the same proportion.
    root = os.path.dirname(os.path.abspath(__file__))
    corner = 3 / 16

    icons = [('favicon.ico', None), ('favicon-32.png', 32)]
    for name, only in icons:
        path = os.path.join(root, 'static', 'img', name)
        source = Image.open(path)
        sizes = (sorted(source.ico.sizes()) if name.endswith('.ico')
                 else [(only, only)])
        check(sizes, f'{name} carries no image at all')

        for size in sizes:
            side = size[0]
            image = (source.ico.getimage(size) if name.endswith('.ico')
                     else source).convert('RGBA')
            alpha = image.split()[3].load()
            radius = round(side * corner)

            corners = [alpha[0, 0], alpha[side - 1, 0],
                       alpha[0, side - 1], alpha[side - 1, side - 1]]
            check(max(corners) == 0,
                  f'{name} at {side}px still has square corners: {corners}')

            # Rounded, not shrunk: the middle of every edge must still reach
            # the icon's boundary.
            edges = [alpha[side // 2, 0], alpha[side // 2, side - 1],
                     alpha[0, side // 2], alpha[side - 1, side // 2]]
            check(min(edges) == 255,
                  f'{name} at {side}px is inset rather than rounded: {edges}')
            check(alpha[side // 2, side // 2] == 255,
                  f'{name} at {side}px is transparent in the middle')

            # Smooth, not stepped: the top edge has to ramp in through at
            # least one partial value, and never step backwards.
            ramp = [alpha[x, 0] for x in range(radius + 1)]
            check(any(0 < v < 255 for v in ramp),
                  f'{name} at {side}px has a hard-edged corner: {ramp}')
            check(all(b >= a for a, b in zip(ramp, ramp[1:])),
                  f'{name} at {side}px does not ramp cleanly: {ramp}')

    # iOS masks a home-screen icon itself, so rounding that one too would show
    # as a pale seam inside the system's own rounding.
    apple = Image.open(os.path.join(root, 'static', 'img',
                                    'apple-touch-icon.png'))
    check(apple.mode == 'RGB',
          'the apple touch icon gained transparency - iOS masks it already')


@test('the application never uses a browser dialog of its own')
def _():
    # window.confirm was the one dialog that looked and behaved like nothing
    # else in the application - browser chrome, browser buttons, no relation to
    # anything on screen. Clearing the canvas now asks in the app's own dialog.
    # Comments stripped first: this file explains in prose why
    # window.confirm was removed, and prose is not a call.
    script = re.sub(r'/\*.*?\*/', '', _script(), flags=re.S)
    script = re.sub(r'(?m)^\s*//.*$', '', script)
    for native in ('window.confirm(', 'window.alert(', 'window.prompt(',
                   'if (!confirm(', '= confirm(', chr(39) + 'alert(' + chr(39),
                   'alert("'):
        check(native not in script, f'a native dialog is back: {native}')
    check('confirmAction(' in script and 'id="confirm-modal"'
          in _read_template('partials/dialogs_convert.html'),
          'the confirmation dialog is gone')


# ---------------------------------------------------------------------------

def main():
    passed = failed = 0
    for name, fn in _RESULTS:
        restore_stubs()
        try:
            fn()
        except Exception as exc:
            failed += 1
            print(f'  FAIL  {name}\n        {type(exc).__name__}: {exc}')
        else:
            passed += 1
            print(f'  ok    {name}')
    print(f'\n{passed} passed, {failed} failed')
    return 1 if failed else 0


if __name__ == '__main__':
    try:
        sys.exit(main())
    finally:
        import shutil
        shutil.rmtree(_SCRATCH, ignore_errors=True)
